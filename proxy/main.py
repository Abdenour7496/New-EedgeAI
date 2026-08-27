"""
Graph-Centric Orchestrated Retrieval (GCOR) proxy — Cognitive Infrastructure.

Architecture
────────────
  Graphiti = knowledge backend (temporal episodes, entity/fact extraction,
             hybrid semantic + keyword + graph + temporal retrieval)
  FalkorDB = storage for Graphiti's graph structure and vector indexes
  MinIO    = durable original-file storage (S3-compatible), keyed by doc_id

This replaced an earlier Neo4j + Qdrant dual-write design — see
docs/adr/0001-graphiti-falkordb-backend.md for the migration rationale.
Confidence filtering, temporal validity, and agent/access-level partitioning
are still applied client-side in _filter_hits() over Graphiti's results.

─────────────────────────────────────────────────────────────────────────────
Retrieval flow for every /v1/chat/completions request
─────────────────────────────────────────────────────
  1. INTENT CLASSIFICATION
     factual | planning | dependency | memory | semantic | inference | belief
     (used for context formatting/logging; Graphiti's search is not
     intent-specific — see build_gcor_context())

  2. HYBRID RETRIEVAL  (Graphiti / FalkorDB)
     graphiti_search() → temporal facts, min-max normalised and filtered by
     score, confidence, temporal validity, and access control (_filter_hits)

  3. REFLECTION CHECK
     No graph records but facts exist → fall back to fact text
     Both empty → LLM general knowledge

  4. CONTEXT INJECTION
     Structured system message including confidence scores

Other endpoints
───────────────
  GET  /v1/models              → live OpenAI model list (with fallback)
  POST /v1/embeddings          → proxied to OpenAI
  POST /v1/buzz/chat/completions → bearer-authenticated GCOR entry point for the
                                    Buzz collaboration bridge (see BUZZ_BRIDGE_API_KEY)
  GET  /health                 → liveness check
"""

import hashlib
import hmac
import io
import json
import logging
import os
import re
import time
import uuid
from datetime import datetime, timezone

import httpx
from fastapi import BackgroundTasks, FastAPI, File, Form, HTTPException, Query, Request, UploadFile
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, StreamingResponse
from prometheus_client import Counter, Histogram
from prometheus_fastapi_instrumentator import Instrumentator
from document_intel import process_document, DocIntelResult
from governance import validate_access_level

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

app = FastAPI()

# ── API authentication ────────────────────────────────────────────────────────
# GCOR_API_KEY is a shared secret every trusted caller (OpenWebUI, its Functions,
# ingest-watcher, admin scripts/curl) must present as `Authorization: Bearer
# <key>`. Without it, /api/* and /v1/* accept requests from anyone who can
# reach this port — see docs/adr/0002-proxy-api-authentication.md.
#
# If GCOR_API_KEY is unset, auth is left OPEN (dev convenience, matches this
# repo's pattern of degrading gracefully rather than crashing on missing
# config) but it's loud about it so this can't be silently shipped to
# production.
GCOR_API_KEY = os.getenv("GCOR_API_KEY", "")
if not GCOR_API_KEY:
    logger.warning(
        "GCOR_API_KEY is not set — /api/* and /v1/* endpoints are UNAUTHENTICATED. "
        "Set GCOR_API_KEY before exposing this stack beyond a trusted dev network."
    )

# Paths that stay open even with GCOR_API_KEY configured: static HTML pages
# (they carry no data themselves) and liveness/metrics probes.
_AUTH_EXEMPT_PATHS = {"/knowledge", "/openclaw", "/health", "/metrics", "/docs", "/openapi.json"}


@app.middleware("http")
async def _require_api_key(request: Request, call_next):
    path = request.url.path
    needs_auth = (path.startswith("/api/") or path.startswith("/v1/")) and path not in _AUTH_EXEMPT_PATHS
    if needs_auth and GCOR_API_KEY:
        header = request.headers.get("authorization", "")
        scheme, _, token = header.partition(" ")
        if scheme.lower() != "bearer" or not hmac.compare_digest(token, GCOR_API_KEY):
            return JSONResponse(status_code=401, content={"detail": "Missing or invalid API key"})
    return await call_next(request)


# ── Prometheus metrics ─────────────────────────────────────────────────────────

# Auto-instruments HTTP request count + latency for every FastAPI endpoint
Instrumentator().instrument(app).expose(app)

# GCOR RAG pipeline
METRIC_RAG_REQUESTS = Counter(
    "gcor_rag_requests_total",
    "Total GCOR RAG pipeline invocations",
    ["intent", "fallback_mode"],
)
METRIC_RAG_DURATION = Histogram(
    "gcor_rag_duration_seconds",
    "End-to-end Graphiti retrieval and context-build latency",
    buckets=[0.1, 0.25, 0.5, 1.0, 2.0, 5.0, 10.0, 30.0],
)
METRIC_GRAPHITI_FACTS = Histogram(
    "gcor_graphiti_facts",
    "Number of Graphiti facts returned per RAG request",
    buckets=[0, 1, 2, 3, 4, 5, 6, 8, 10, 15, 20],
)
METRIC_CONTEXT_RECORDS = Histogram(
    "gcor_context_records",
    "Number of knowledge records added to the LLM context",
    buckets=[0, 1, 2, 5, 10, 15, 20, 30, 50],
)

# Embedding
METRIC_EMBED_DURATION = Histogram(
    "gcor_embed_duration_seconds",
    "OpenAI embedding API call latency",
    buckets=[0.05, 0.1, 0.2, 0.5, 1.0, 2.0, 5.0],
)

# LLM calls
METRIC_LLM_REQUESTS = Counter(
    "gcor_llm_requests_total",
    "Total LLM API calls",
    ["backend", "status"],
)
METRIC_LLM_DURATION = Histogram(
    "gcor_llm_duration_seconds",
    "LLM API call latency (non-streaming only)",
    ["backend"],
    buckets=[0.5, 1.0, 2.0, 5.0, 10.0, 20.0, 40.0, 60.0, 120.0],
)

# Ingest
METRIC_INGEST_TOTAL = Counter(
    "gcor_ingest_total",
    "Total document ingest operations",
    ["status"],
)
METRIC_INGEST_CHUNKS = Histogram(
    "gcor_ingest_chunks",
    "Number of chunks produced per ingested document",
    buckets=[1, 5, 10, 20, 50, 100, 200, 500],
)

# Collection management
METRIC_COLLECTION_OPS = Counter(
    "gcor_collection_ops_total",
    "Graphiti group management operations",
    ["operation"],
)

# ── LLM backends ───────────────────────────────────────────────────────────────
OPENAI_API_KEY       = os.getenv("OPENAI_API_KEY", "")
OPENAI_CHAT_MODEL    = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o")
ANTHROPIC_API_KEY    = os.getenv("ANTHROPIC_API_KEY", "")
ANTHROPIC_CHAT_MODEL = os.getenv("ANTHROPIC_CHAT_MODEL", "claude-sonnet-4-6")
LLM_BACKEND          = os.getenv("LLM_BACKEND", "openai")   # openai | anthropic | ollama | openclaw | copilot
VISION_BACKEND       = os.getenv("VISION_BACKEND") or LLM_BACKEND   # separate override for vision calls

# ── GitHub Copilot ──────────────────────────────────────────────────────────────
COPILOT_API_KEY    = os.getenv("COPILOT_API_KEY", "")
COPILOT_BASE_URL   = os.getenv("COPILOT_BASE_URL", "https://api.githubcopilot.com")
COPILOT_CHAT_MODEL = os.getenv("COPILOT_CHAT_MODEL", "gpt-4.1")

# ── OpenClaw ────────────────────────────────────────────────────────────────────
OPENCLAW_BASE_URL      = os.getenv("OPENCLAW_BASE_URL",      "http://openclaw:18799/v1")
OPENCLAW_GATEWAY_TOKEN = os.getenv("OPENCLAW_GATEWAY_TOKEN", "")

# ── Ollama ─────────────────────────────────────────────────────────────────────
OLLAMA_BASE_URL        = os.getenv("OLLAMA_BASE_URL",        "http://ollama:11434")
OLLAMA_MODEL           = os.getenv("OLLAMA_MODEL",           "llama3.2")
OLLAMA_EMBEDDING_MODEL = os.getenv("OLLAMA_EMBEDDING_MODEL", "nomic-embed-text")

# ── Embedding ──────────────────────────────────────────────────────────────────
EMBEDDING_BACKEND = os.getenv("EMBEDDING_BACKEND", "openai")   # openai | ollama
EMBEDDING_MODEL   = os.getenv("EMBEDDING_MODEL",   "text-embedding-3-small")
EMBED_BATCH_SIZE  = int(os.getenv("EMBED_BATCH_SIZE",  "32"))   # chunks per OpenAI call
EMBED_BATCH_DELAY = float(os.getenv("EMBED_BATCH_DELAY", "0.5")) # seconds between batches
EMBED_MAX_RETRIES = int(os.getenv("EMBED_MAX_RETRIES",  "6"))    # 429/5xx retry attempts
LLM_MAX_RETRIES   = int(os.getenv("LLM_MAX_RETRIES",    "3"))    # 429/529/5xx chat retry attempts
LLM_RETRY_BASE    = float(os.getenv("LLM_RETRY_BASE",  "1.0"))   # base delay for LLM retries

# ── Graphiti temporal graph (FalkorDB backend) ───────────────────────────────
GRAPHITI_URL      = os.getenv("GRAPHITI_URL", "http://graphiti:8000").rstrip("/")
GRAPHITI_GROUP_ID = os.getenv("GRAPHITI_GROUP_ID", "documents")
GRAPHITI_TOP_K    = int(os.getenv("GRAPHITI_TOP_K", "8"))
# Local CPU-only extraction (the default qwen2.5:7b via Ollama) can take
# several minutes per ingest call, especially with multiple chunks — each
# chunk is a sequential LLM extraction inside Graphiti. 300s was observed to
# be too short even for a 2-message chat session on CPU-only hardware, which
# surfaced as a 502 to the caller despite the episode landing successfully a
# short time later. A hosted OpenAI-compatible GRAPHITI_API_BASE_URL will
# need far less of this margin.
GRAPHITI_INGEST_TIMEOUT_SECONDS = float(os.getenv("GRAPHITI_INGEST_TIMEOUT_SECONDS", "1800"))

# Compatibility aliases retained for the existing Knowledge UI route names.
QDRANT_COLLECTION = GRAPHITI_GROUP_ID
QDRANT_TOP_K      = GRAPHITI_TOP_K

# Cached embedding dimension — probed once at first use
_embed_dim_cache: int | None = None


async def _get_embed_dim() -> int:
    """Return the vector dimension of the current embedding model (cached)."""
    global _embed_dim_cache
    if _embed_dim_cache is not None:
        return _embed_dim_cache
    vec = await embed_text("dimension probe")
    _embed_dim_cache = len(vec) if vec else 768
    return _embed_dim_cache

ENABLE_RAG = os.getenv("ENABLE_RAG", "true").lower() in ("true", "1", "yes")


async def graphiti_search(query: str, group_id: str | None = None) -> tuple[list, list]:
    """Search Graphiti and adapt temporal facts to the proxy's context contract."""
    group = group_id or GRAPHITI_GROUP_ID
    async with httpx.AsyncClient(timeout=30) as client:
        response = await client.post(
            f"{GRAPHITI_URL}/search",
            json={"group_ids": [group], "query": query, "max_facts": GRAPHITI_TOP_K},
        )
        response.raise_for_status()
    facts = response.json().get("facts", [])
    hits = []
    records = []
    for rank, fact in enumerate(facts):
        payload = {
            "text": fact.get("fact", ""),
            "graphiti_edge_uuid": fact.get("uuid", ""),
            "document_id": (fact.get("episodes") or [""])[0],
            "document_title": fact.get("name", "Graphiti fact"),
            "confidence": 1.0,
            "valid_from": fact.get("valid_at") or fact.get("created_at"),
            "valid_to": fact.get("invalid_at") or fact.get("expired_at"),
            "access_level": "public",
        }
        hits.append({"id": fact.get("uuid"), "score": 1.0 - rank / max(len(facts), 1), "payload": payload})
        records.append({
            "node": payload,
            "labels": ["TemporalFact"],
            "document": {"title": fact.get("name", "Graphiti fact")},
            "related": [],
        })
    return hits, records


async def graphiti_ingest(
    group_id: str, doc_id: str, title: str, source: str, chunks: list[str],
    agent_id: str, access_level: str, valid_to: str | None = None,
) -> int:
    """Queue document chunks as Graphiti episodes in a group namespace."""
    now = datetime.now(timezone.utc).isoformat()
    messages = []
    for position, chunk in enumerate(chunks):
        metadata = (
            f"Document: {title}\nSource: {source}\nDocument ID: {doc_id}\n"
            f"Chunk: {position + 1}/{len(chunks)}\nAgent: {agent_id or 'shared'}\n"
            f"Access: {access_level}\nValid until: {valid_to or 'open'}\n\n"
        )
        messages.append({
            "name": f"{title} — chunk {position + 1}",
            "content": metadata + chunk,
            "role_type": "system",
            "role": "document",
            "timestamp": now,
            "source_description": source,
        })
    async with httpx.AsyncClient(timeout=GRAPHITI_INGEST_TIMEOUT_SECONDS) as client:
        response = await client.post(
            f"{GRAPHITI_URL}/messages", json={"group_id": group_id, "messages": messages}
        )
        response.raise_for_status()
    return len(messages)

# ── Cognitive infrastructure knobs ────────────────────────────────────────────
# Minimum confidence (0.0–1.0) — hits below this are discarded
CONFIDENCE_THRESHOLD = float(os.getenv("CONFIDENCE_THRESHOLD", "0.0"))
# Agent id scopes memory/belief/inference retrieval to a single partition
AGENT_ID = os.getenv("AGENT_ID", "")

# ── Buzz bridge (optional --profile buzz integration) ─────────────────────────
# Shared secret the buzz-agent-bridge container presents on /v1/buzz/chat/completions.
# That route is reachable from the Buzz relay's channel traffic (humans + agents),
# a less-trusted source than the internal-network-only /v1/chat/completions, so it
# is closed (503) unless this is explicitly set.
BUZZ_BRIDGE_API_KEY = os.getenv("BUZZ_BRIDGE_API_KEY", "")
# Optional allowlist restricting which collections Buzz-originated queries may target.
# Empty = no restriction beyond what /v1/chat/completions already allows.
BUZZ_ALLOWED_COLLECTIONS = [
    c.strip() for c in os.getenv("BUZZ_ALLOWED_COLLECTIONS", "").split(",") if c.strip()
]
BUZZ_BRIDGE_RATE_LIMIT = int(os.getenv("BUZZ_BRIDGE_RATE_LIMIT_PER_MIN", "30"))
# buzz-acp bundles base prompt + up to BUZZ_ACP_CONTEXT_MESSAGE_LIMIT prior
# messages + memory injection into a single message, easily tens of KB —
# 8000 chars (initial guess) was too tight for real traffic and rejected
# legitimate turns outright.
BUZZ_BRIDGE_MAX_MESSAGE_CHARS = int(os.getenv("BUZZ_BRIDGE_MAX_MESSAGE_CHARS", "100000"))
# Cap on /v1/buzz/ingest uploads — this route is reachable from Buzz channel
# content (an image attachment or a storage_key lookup), so unlike the
# internal-network-only /api/ingest it needs its own size ceiling.
BUZZ_BRIDGE_MAX_UPLOAD_BYTES = int(os.getenv("BUZZ_BRIDGE_MAX_UPLOAD_BYTES", str(20 * 1024 * 1024)))

# ── Document storage (MinIO, S3-compatible) ───────────────────────────────────
# Best-effort: ingestion still works if MinIO is unreachable or unconfigured,
# it just skips persisting the original file.
S3_ENDPOINT_URL      = os.getenv("S3_ENDPOINT_URL", "http://minio:9000")
S3_ACCESS_KEY        = os.getenv("S3_ACCESS_KEY", "")
S3_SECRET_KEY        = os.getenv("S3_SECRET_KEY", "")
S3_BUCKET            = os.getenv("S3_BUCKET", "documents")
S3_ORIGINALS_PREFIX  = os.getenv("S3_ORIGINALS_PREFIX", "originals/")

_s3_client = None
_s3_client_init = False


def _get_s3_client():
    """Lazily construct the boto3 S3 client (MinIO). Returns None if unconfigured."""
    global _s3_client, _s3_client_init
    if not _s3_client_init:
        _s3_client_init = True
        if S3_ACCESS_KEY and S3_SECRET_KEY:
            import boto3
            from botocore.client import Config as BotoConfig
            _s3_client = boto3.client(
                "s3",
                endpoint_url=S3_ENDPOINT_URL,
                aws_access_key_id=S3_ACCESS_KEY,
                aws_secret_access_key=S3_SECRET_KEY,
                config=BotoConfig(signature_version="s3v4"),
                region_name="us-east-1",
            )
    return _s3_client


_BUCKET_NAME_RE = re.compile(r"^[a-z0-9][a-z0-9\-]{1,61}[a-z0-9]$")

# Buckets we've already confirmed exist this process — avoids a round-trip
# per ingest once a bucket has been seen.
_known_buckets: set[str] = set()


def _validate_bucket_name(name: str) -> str:
    """S3/MinIO bucket naming rules: 3-63 chars, lowercase alphanumeric + hyphens,
    no leading/trailing hyphen. Raises HTTPException on failure."""
    if not _BUCKET_NAME_RE.match(name):
        raise HTTPException(
            status_code=400,
            detail=(
                f"Invalid bucket name '{name}': must be 3-63 characters, "
                "lowercase letters, numbers and hyphens only, no leading/trailing hyphen"
            ),
        )
    return name


async def _ensure_bucket_exists(client, bucket: str) -> None:
    """Create the bucket if it doesn't already exist. No-op once seen."""
    if bucket in _known_buckets:
        return
    try:
        await asyncio.to_thread(client.head_bucket, Bucket=bucket)
    except Exception:
        try:
            await asyncio.to_thread(client.create_bucket, Bucket=bucket)
        except Exception as exc:
            # MinIO/S3 both raise a (different) "already exists/owned by you"
            # error on a race here — anything else is a real failure.
            if "BucketAlready" not in str(exc):
                raise
    _known_buckets.add(bucket)


async def _store_original_to_s3(
    doc_id: str, filename: str, data: bytes, bucket: str = S3_BUCKET,
) -> str | None:
    """Persist the raw uploaded file to MinIO, creating `bucket` if needed.
    Returns the object key, or None on failure/unconfigured."""
    client = _get_s3_client()
    if client is None:
        return None
    key = f"{S3_ORIGINALS_PREFIX}{doc_id}/{filename}"
    try:
        await _ensure_bucket_exists(client, bucket)
        await asyncio.to_thread(client.put_object, Bucket=bucket, Key=key, Body=data)
        return key
    except Exception as exc:
        logger.warning("S3 store of original file failed (doc_id=%s, bucket=%s): %s", doc_id, bucket, exc)
        return None


async def _find_and_fetch_from_s3(name: str) -> tuple[bytes | None, str | None]:
    """Look up a file by bare name or key across the documents bucket's known
    locations: exact key, inbox/, processed/, failed/, and originals/<doc_id>/
    (doc_id unknown, so that one is a suffix match). Returns (bytes, filename)
    or (None, None) if not found or S3 isn't configured."""
    client = _get_s3_client()
    if client is None:
        return None, None

    def _try_get(key: str):
        try:
            return client.get_object(Bucket=S3_BUCKET, Key=key)["Body"].read()
        except Exception:
            return None

    for prefix in ("", "inbox/", "processed/", "failed/"):
        data = await asyncio.to_thread(_try_get, f"{prefix}{name}")
        if data is not None:
            return data, name.rsplit("/", 1)[-1]

    def _search_originals() -> str | None:
        paginator = client.get_paginator("list_objects_v2")
        for page in paginator.paginate(Bucket=S3_BUCKET, Prefix=S3_ORIGINALS_PREFIX):
            for obj in page.get("Contents", []):
                if obj["Key"].endswith(f"/{name}"):
                    return obj["Key"]
        return None

    match_key = await asyncio.to_thread(_search_originals)
    if match_key:
        data = await asyncio.to_thread(_try_get, match_key)
        if data is not None:
            return data, name

    return None, None


# ── Step 1: Intent classification ─────────────────────────────────────────────

_INTENT_KEYWORDS = {
    "planning":    ["plan", "schedule", "next step", "roadmap", "workflow",
                    "milestone", "sprint", "agenda", "timeline", "sequence"],
    "dependency":  ["depends", "depend on", "requires", "blocked", "prerequisite", "constraint",
                    "relies on", "linked to", "what does", "what do"],
    "memory":      ["remember", "recall", "history", "last time", "previously",
                    "past", "earlier", "before we", "you said", "i told"],
    "factual":     ["what is", "how does", "why", "explain", "define", "describe",
                    "tell me about", "how do", "what are"],
    "inference":   ["infer", "conclude", "deduce", "reasoning", "derive",
                    "implication", "therefore", "follows that", "suggests",
                    "implies", "logical", "because of"],
    "belief":      ["believe", "think", "opinion", "doubt", "uncertain",
                    "probably", "likely", "assume", "suppose", "suspect",
                    "my view", "i think"],
}

# ── Embedding-based intent classification ─────────────────────────────────────

_INTENT_EXEMPLARS = {
    "factual":    "What is this? How does it work? Explain and define the concept.",
    "planning":   "What are the next steps? Create a roadmap and schedule milestones.",
    "dependency": "What does this depend on? What blocks progress? What are prerequisites?",
    "memory":     "What did we discuss before? Recall our previous conversation history.",
    "inference":  "What can we conclude? Derive the logical implication from this evidence.",
    "belief":     "What do you think? What is your opinion? What do you believe is likely?",
    "semantic":   "Search for relevant information about this topic.",
}

_INTENT_EXEMPLAR_VECS: dict = {}


def _cosine_sim(a: list, b: list) -> float:
    """Cosine similarity between two equal-length vectors."""
    if not a or not b or len(a) != len(b):
        return 0.0
    dot = sum(x * y for x, y in zip(a, b))
    norm_a = sum(x * x for x in a) ** 0.5
    norm_b = sum(x * x for x in b) ** 0.5
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


@app.on_event("startup")
async def _precompute_intent_exemplars():
    """Embed intent exemplars at startup for embedding-based classification."""
    global _INTENT_EXEMPLAR_VECS
    try:
        for intent, sentence in _INTENT_EXEMPLARS.items():
            vec = await embed_text(sentence)
            if vec:
                _INTENT_EXEMPLAR_VECS[intent] = vec
        logger.info("Intent exemplar embeddings precomputed: %d intents", len(_INTENT_EXEMPLAR_VECS))
    except Exception as exc:
        logger.warning("Could not precompute intent exemplars (will use keyword fallback): %s", exc)


async def classify_intent(query: str) -> str:
    """Classify query intent via embedding cosine sim; falls back to keyword scan."""
    # Embedding-based path
    if _INTENT_EXEMPLAR_VECS:
        try:
            query_vec = await embed_text(query)
            if query_vec:
                scores = {
                    intent: _cosine_sim(query_vec, vec)
                    for intent, vec in _INTENT_EXEMPLAR_VECS.items()
                }
                best = max(scores, key=scores.get)
                logger.debug("Intent scores: %s → %s", scores, best)
                return best
        except Exception as exc:
            logger.warning("Embedding intent classification failed, using keyword fallback: %s", exc)
    # Keyword fallback
    q = query.lower()
    for intent, keywords in _INTENT_KEYWORDS.items():
        if any(kw in q for kw in keywords):
            return intent
    return "semantic"


# ── Step 2a: Semantic phase (Qdrant) ──────────────────────────────────────────

def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


import asyncio
import random

async def _embed_with_retry(
    payload: dict,
    *,
    max_retries: int = EMBED_MAX_RETRIES,
) -> dict:
    """
    POST to OpenAI /v1/embeddings with exponential-backoff retry on 429 / 5xx.

    Respects the Retry-After header when present.
    Raises httpx.HTTPStatusError on permanent failure.
    """
    base_delay = 1.0
    for attempt in range(max_retries + 1):
        async with httpx.AsyncClient(timeout=60) as c:
            resp = await c.post(
                "https://api.openai.com/v1/embeddings",
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}",
                         "Content-Type": "application/json"},
                json=payload,
            )

        if resp.status_code == 200:
            return resp.json()

        retry_after = None
        if resp.status_code == 429 or resp.status_code >= 500:
            # Honour Retry-After header if present
            ra = resp.headers.get("retry-after") or resp.headers.get("x-ratelimit-reset-requests")
            if ra:
                try:
                    retry_after = float(ra)
                except ValueError:
                    pass

            if attempt < max_retries:
                wait = retry_after if retry_after else (base_delay * (2 ** attempt) + random.uniform(0, 0.5))
                logger.warning(
                    "OpenAI embeddings %s (attempt %d/%d) — waiting %.1fs",
                    resp.status_code, attempt + 1, max_retries, wait,
                )
                await asyncio.sleep(wait)
                continue

        resp.raise_for_status()   # permanent error

    resp.raise_for_status()


async def _embed_batch(texts: list[str]) -> list[list[float]]:
    """Embed a batch of texts via the configured backend (OpenAI or Ollama)."""
    if EMBEDDING_BACKEND == "ollama":
        async with httpx.AsyncClient(timeout=120) as c:
            resp = await c.post(
                f"{OLLAMA_BASE_URL}/v1/embeddings",
                json={"input": texts, "model": OLLAMA_EMBEDDING_MODEL},
            )
            resp.raise_for_status()
        items = resp.json()["data"]
        items.sort(key=lambda x: x["index"])
        return [d["embedding"] for d in items]
    data = await _embed_with_retry({"input": texts, "model": EMBEDDING_MODEL})
    items = data["data"]
    items.sort(key=lambda x: x["index"])
    return [d["embedding"] for d in items]


async def embed_text(text: str) -> list | None:
    try:
        t0 = time.monotonic()
        if EMBEDDING_BACKEND == "ollama":
            async with httpx.AsyncClient(timeout=60) as c:
                resp = await c.post(
                    f"{OLLAMA_BASE_URL}/v1/embeddings",
                    json={"input": text[:8000], "model": OLLAMA_EMBEDDING_MODEL},
                )
                resp.raise_for_status()
            vec = resp.json()["data"][0]["embedding"]
        else:
            if not OPENAI_API_KEY:
                return None
            data = await _embed_with_retry({"input": text[:8000], "model": EMBEDDING_MODEL})
            vec = data["data"][0]["embedding"]
        METRIC_EMBED_DURATION.observe(time.monotonic() - t0)
        return vec
    except Exception as exc:
        logger.warning("Embedding failed: %s", exc)
        return None


# Collections to search during RAG (in addition to QDRANT_COLLECTION).
# Only collections whose dimension matches the active embedding model are queried.
RAG_EXTRA_COLLECTIONS = [
    c.strip()
    for c in os.getenv("RAG_EXTRA_COLLECTIONS", "buddy_memory").split(",")
    if c.strip()
]


def _request_rag_collection(body: dict, request: Request) -> str | None:
    """Return an optional collection selected by an OpenWebUI workspace filter."""
    body_collection = body.pop("gcor_collection", None)
    requested = request.headers.get("X-GCOR-Collection") or body_collection
    if requested is None:
        return None
    if not isinstance(requested, str):
        raise HTTPException(status_code=400, detail="gcor_collection must be a string")

    collection = requested.strip()
    if not collection:
        return None
    if len(collection) > 255 or any(char in collection for char in "/\\?#"):
        raise HTTPException(status_code=400, detail="Invalid gcor_collection")
    return collection


def _filter_hits(hits: list) -> list:
    """
    Apply cognitive filters to raw Qdrant hits:
      1. Confidence threshold — payload.confidence below minimum is discarded
      2. Temporal validity   — expired valid_to is discarded
      3. Access control      — access_level incompatible with AGENT_ID is discarded
    """
    now = _now_iso()
    filtered = []
    for h in hits:
        p = h.get("payload") or {}

        # ── confidence ────────────────────────────────────────────────────────
        node_confidence = p.get("confidence")
        if node_confidence is not None and node_confidence < CONFIDENCE_THRESHOLD:
            logger.debug("Dropping hit %s: confidence %.2f < %.2f",
                         h.get("id"), node_confidence, CONFIDENCE_THRESHOLD)
            continue

        # ── temporal validity ─────────────────────────────────────────────────
        valid_to = p.get("valid_to")
        if valid_to and valid_to < now:
            logger.debug("Dropping hit %s: expired valid_to %s", h.get("id"), valid_to)
            continue
        valid_from = p.get("valid_from")
        if valid_from and valid_from > now:
            logger.debug("Dropping hit %s: not yet valid (valid_from %s)", h.get("id"), valid_from)
            continue

        # ── access control ────────────────────────────────────────────────────
        # `or "public"` (not just a dict default) because a payload can carry
        # access_level=None explicitly, not merely omit the key.
        access_level = p.get("access_level") or "public"
        if access_level == "restricted":
            logger.debug("Dropping hit %s: restricted access", h.get("id"))
            continue
        if access_level.startswith("agent:") and AGENT_ID:
            owner = access_level.split(":", 1)[1]
            if owner != AGENT_ID:
                logger.debug("Dropping hit %s: owned by %s, not %s", h.get("id"), owner, AGENT_ID)
                continue

        filtered.append(h)
    return filtered


# ── Step 2b/2c: structural expansion + provenance chain ──────────────────────
# The Neo4j-based graph expansion (_EXPAND_CYPHER/neo4j_expand) and provenance
# chain (build_provenance_chain) that used to run here were removed with the
# Graphiti/FalkorDB cutover — see docs/adr/0001-graphiti-falkordb-backend.md.
# graphiti_search() above already returns hybrid semantic + graph + temporal
# results in one call, and _run_chat_completion() builds its provenance string
# directly from each hit's document_id, so neither helper had a live caller.


# ── Step 3: Reflection ─────────────────────────────────────────────────────────

def _reflection_fallback(qdrant_hits: list, graph_records: list) -> list:
    if graph_records:
        return graph_records
    if qdrant_hits:
        logger.info("Reflection: no graph records, falling back to Graphiti fact text.")
        return [
            {
                "node": {
                    "text":        h["payload"].get("text", ""),
                    "document_id": h["payload"].get("document_id", ""),
                    "confidence":  h["payload"].get("confidence"),
                    "valid_from":  h["payload"].get("valid_from"),
                    "valid_to":    h["payload"].get("valid_to"),
                },
                "labels":   [h["payload"].get("node_type", "Chunk")],
                "document": {"title": h["payload"].get("document_title", "")},
            }
            for h in qdrant_hits
            if h.get("payload", {}).get("text")
        ]
    return []


# ── Context builder ────────────────────────────────────────────────────────────

def _confidence_badge(props: dict) -> str:
    c = props.get("confidence")
    if c is None:
        return ""
    pct = int(c * 100)
    if pct >= 90:
        tier = "high"
    elif pct >= 60:
        tier = "medium"
    else:
        tier = "low"
    return f"  [confidence: {pct}% / {tier}]"


def _temporal_badge(props: dict) -> str:
    now = _now_iso()[:10]
    parts = []
    vf = props.get("valid_from")
    vt = props.get("valid_to")
    if vf:
        vf_short = vf[:10]
        parts.append(f"from {vf_short}")
        # Add staleness hint for old knowledge
        try:
            from datetime import datetime, timezone
            age_days = (datetime.now(timezone.utc) - datetime.fromisoformat(vf)).days
            if age_days > 365:
                parts.append(f"⚠ {age_days // 365}y old")
            elif age_days > 90:
                parts.append(f"⚠ {age_days}d old")
        except Exception:
            pass
    if vt:
        vt_short = vt[:10]
        parts.append(f"expires {vt_short}")
        # Warn if expiring soon
        if vt_short <= now:
            parts.append("EXPIRED")
        elif vt[:10] <= (now[:8] + str(int(now[8:10]) + 7).zfill(2))[:10]:
            parts.append("expiring soon")
    return f"  [{', '.join(parts)}]" if parts else ""


def resolve_belief_conflicts(records: list) -> list:
    """
    For records with belief contradictions, pick the winning belief by:
    1. Highest confidence
    2. Most recent valid_from as tiebreaker
    Sets _resolved_belief and _resolution_reason on each conflicting record.
    """
    for rec in records:
        contradictions = [c for c in (rec.get("contradictions") or []) if c]
        if not contradictions:
            continue
        node = rec.get("node", {})
        candidates = [node] + contradictions
        def _sort_key(b):
            conf = float(b.get("confidence") or 0)
            vf   = str(b.get("valid_from") or "")
            return (conf, vf)
        winner = max(candidates, key=_sort_key)
        conf_pct = int(float(winner.get("confidence") or 0) * 100)
        vf = str(winner.get("valid_from") or "")[:10]
        rec["_resolved_belief"] = winner
        rec["_resolution_reason"] = (
            f"Selected belief with confidence={conf_pct}%"
            + (f", valid_from={vf}" if vf else "")
            + f" (resolved from {len(candidates)} candidates)"
        )
    return records


def build_gcor_context(intent: str, qdrant_hits: list, graph_records: list, provenance: str = "") -> str:
    records = _reflection_fallback(qdrant_hits, graph_records)
    if not records:
        return ""

    confidence_note = (
        f"confidence ≥ {int(CONFIDENCE_THRESHOLD*100)}% "
        if CONFIDENCE_THRESHOLD > 0 else ""
    )
    agent_note = f"agent partition: {AGENT_ID} " if AGENT_ID else ""

    now = _now_iso()
    lines = [
        f"## Retrieved Context  [intent: {intent} | "
        f"{len(qdrant_hits)} Graphiti facts → {len(records)} graph records"
        + (f" | {confidence_note}{agent_note}".rstrip() if (confidence_note or agent_note) else "")
        + "]",
        "",
        f"Current time: {now[:19]}Z. "
        "Graphiti on FalkorDB is the source of truth. Results combine semantic, keyword, "
        "graph, and temporal retrieval. Prefer currently valid and more recent facts "
        "when information conflicts.",
    ]
    if provenance:
        lines += ["", f"**Retrieval Provenance:** {provenance}"]
    lines.append("")

    for rec in records:
        node   = rec.get("node") or {}
        labels = rec.get("labels") or []
        doc    = rec.get("document") or {}

        node_type = labels[0] if labels else "Node"
        content   = (
            node.get("text") or node.get("content") or
            node.get("description") or node.get("title") or
            node.get("subject") or str(node)
        )[:600]

        conf_badge    = _confidence_badge(node)
        temporal_badge = _temporal_badge(node)
        lines.append(f"### [{node_type}]{conf_badge}{temporal_badge}")
        lines.append(content)

        if doc.get("title") or doc.get("source"):
            lines.append(f"Source: {doc.get('title') or doc.get('source', '')}")

        # Reasoning trace for Inference nodes
        if node.get("reasoning_trace"):
            lines.append(f"Reasoning: {node['reasoning_trace'][:300]}")

        # Intent-specific extras
        if intent == "planning":
            deps = [d for d in (rec.get("dependencies") or []) if d]
            if deps:
                lines.append("Dependencies: " + " → ".join(
                    d.get("description") or d.get("title") or str(d) for d in deps
                ))
            blocking = [g for g in (rec.get("blocking_goals") or []) if g]
            if blocking:
                lines.append("Blocked by: " + ", ".join(
                    g.get("description") or str(g) for g in blocking
                ))
            inferences = [i for i in (rec.get("supporting_inferences") or []) if i]
            if inferences:
                lines.append("Supporting inferences: " + "; ".join(
                    f"{i.get('text', '')[:80]} ({int(i.get('confidence', 0)*100)}%)"
                    for i in inferences
                ))

        elif intent == "dependency":
            chain = [x for x in (rec.get("dependency_chain") or []) if x]
            if len(chain) > 1:
                lines.append("Dependency path: " + " → ".join(
                    x.get("description") or x.get("title") or str(x) for x in chain
                ))

        elif intent == "memory":
            mems = [m for m in (rec.get("memories") or []) if m]
            if mems:
                lines.append("Memories: " + "; ".join(
                    f"{m.get('content', '')[:80]}{_confidence_badge(m)}"
                    for m in mems[:3]
                ))
            beliefs = [b for b in (rec.get("beliefs") or []) if b]
            if beliefs:
                lines.append("Beliefs: " + "; ".join(
                    f"{b.get('content', '')[:80]}{_confidence_badge(b)}"
                    for b in beliefs[:3]
                ))
            evts = sorted(
                [e for e in (rec.get("events") or []) if e],
                key=lambda event: str(event.get("timestamp") or ""),
                reverse=True,
            )
            if evts:
                lines.append("Events: " + "; ".join(
                    f"{e.get('timestamp', '')[:10]} {e.get('description', '')}"
                    for e in evts[:3]
                ))

        elif intent == "inference":
            srcs = [s for s in (rec.get("sources") or []) if s]
            if srcs:
                lines.append("Derived from: " + "; ".join(
                    (s.get("text") or s.get("content") or str(s))[:80]
                    for s in srcs[:3]
                ))
            supports = [t for t in (rec.get("supports") or []) if t]
            if supports:
                lines.append("Supports: " + "; ".join(
                    (t.get("description") or t.get("content") or str(t))[:80]
                    for t in supports[:3]
                ))
            downstream = [d for d in (rec.get("downstream_inferences") or []) if d]
            if downstream:
                lines.append("Downstream inferences: " + "; ".join(
                    f"{d.get('text', '')[:60]}{_confidence_badge(d)}"
                    for d in downstream[:2]
                ))

        elif intent == "belief":
            contradictions = [c for c in (rec.get("contradictions") or []) if c]
            if contradictions:
                lines.append("Contradicts: " + "; ".join(
                    f"{c.get('content', '')[:80]}{_confidence_badge(c)}"
                    for c in contradictions[:2]
                ))
            inferences = [i for i in (rec.get("supporting_inferences") or []) if i]
            if inferences:
                lines.append("Supported by: " + "; ".join(
                    f"{i.get('text', '')[:60]}{_confidence_badge(i)}"
                    for i in inferences[:2]
                ))
            agent = rec.get("agent") or {}
            if agent.get("id"):
                lines.append(f"Held by agent: {agent['id']}")
            if rec.get("_resolution_reason"):
                lines.append(f"⚖️ Conflict resolution: {rec['_resolution_reason']}")

        elif intent == "semantic":
            concepts = [c for c in (rec.get("concepts") or []) if c]
            if concepts:
                lines.append("Concepts: " + ", ".join(c.get("name", "") for c in concepts))
            neighbors = [n for n in (rec.get("neighbors") or []) if n]
            if neighbors:
                lines.append("Adjacent chunks: " + " | ".join(
                    n.get("text", "")[:100] for n in neighbors
                ))

        else:  # factual
            related = [r for r in (rec.get("related") or []) if r and r.get("props")]
            if related:
                lines.append("Related: " + "; ".join(
                    f"{r['rel']} → {r['props'].get('name') or r['props'].get('text', '')[:80]}"
                    for r in related[:4]
                ))

        lines.append("")

    lines.append(
        "Instruction: use the retrieved context above when relevant. "
        "Respect the confidence scores — lower-confidence information should be "
        "presented with appropriate uncertainty. If the context is insufficient, "
        "apply general knowledge."
    )
    return "\n".join(lines)


def inject_context(messages: list, context: str) -> list:
    if not context:
        return messages
    result, merged = [], False
    for m in messages:
        if m.get("role") == "system" and not merged:
            result.append({"role": "system",
                           "content": context + "\n\n---\n\n" + m.get("content", "")})
            merged = True
        else:
            result.append(m)
    if not merged:
        result = [{"role": "system", "content": context}] + messages
    return result


def last_user_text(messages: list) -> str:
    for m in reversed(messages):
        if m.get("role") == "user":
            content = m.get("content", "")
            if isinstance(content, str):
                return content
            if isinstance(content, list):
                return " ".join(
                    p.get("text", "") for p in content
                    if isinstance(p, dict) and p.get("type") == "text"
                )
    return ""


# ── LLM routing ───────────────────────────────────────────────────────────────

def _resolve_model(model: str) -> tuple:
    if model == "openclaw":
        if LLM_BACKEND == "openclaw":
            return "openclaw", "openclaw"
        if LLM_BACKEND == "ollama":
            return OLLAMA_MODEL, "ollama"
        if LLM_BACKEND == "openai":
            return OPENAI_CHAT_MODEL, "openai"
        return ANTHROPIC_CHAT_MODEL, "anthropic"
    if model.startswith("claude-"):
        return model, "anthropic"
    if model.startswith(("gpt-", "o1", "o3")):
        return model, "openai"
    if LLM_BACKEND == "openclaw":
        return "openclaw", "openclaw"
    if LLM_BACKEND == "anthropic":
        return ANTHROPIC_CHAT_MODEL, "anthropic"
    if LLM_BACKEND == "ollama":
        return OLLAMA_MODEL, "ollama"
    return OPENAI_CHAT_MODEL, "openai"


async def call_ollama(body: dict):
    """Forward chat completions to the local Ollama OpenAI-compatible endpoint."""
    url = f"{OLLAMA_BASE_URL}/v1/chat/completions"
    body = {**body, "model": body.get("model", OLLAMA_MODEL)}
    if body.get("stream"):
        async def stream():
            try:
                METRIC_LLM_REQUESTS.labels(backend="ollama", status="started").inc()
                async with httpx.AsyncClient(timeout=300) as c:
                    async with c.stream("POST", url, json=body) as resp:
                        if resp.status_code != 200:
                            err = await resp.aread()
                            logger.error("Ollama %s: %s", resp.status_code, err[:200])
                            METRIC_LLM_REQUESTS.labels(backend="ollama", status="error").inc()
                            yield f'data: {{"error": "Ollama {resp.status_code}"}}\n\n'.encode()
                            return
                        METRIC_LLM_REQUESTS.labels(backend="ollama", status="success").inc()
                        async for chunk in resp.aiter_bytes():
                            yield chunk
            except Exception as exc:
                logger.error("Ollama stream error: %s", exc)
                METRIC_LLM_REQUESTS.labels(backend="ollama", status="exception").inc()
                yield b'data: {"error": "proxy error"}\n\n'
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        try:
            async with httpx.AsyncClient(timeout=300) as c:
                resp = await c.post(url, json=body)
                resp.raise_for_status()
            METRIC_LLM_REQUESTS.labels(backend="ollama", status="success").inc()
            METRIC_LLM_DURATION.labels(backend="ollama").observe(time.monotonic() - t0)
            return JSONResponse(content=resp.json(), status_code=resp.status_code)
        except Exception:
            METRIC_LLM_REQUESTS.labels(backend="ollama", status="error").inc()
            raise


async def call_openclaw(body: dict):
    """Forward chat completions to the openclaw gateway OpenAI-compatible endpoint."""
    url = f"{OPENCLAW_BASE_URL}/chat/completions"
    headers = {"Authorization": f"Bearer {OPENCLAW_GATEWAY_TOKEN}", "Content-Type": "application/json"}
    body = {**body, "model": body.get("model", "default")}
    if body.get("stream"):
        async def normalized_openclaw_stream(resp):
            terminal_chunk_sent = False
            chunk_meta = {
                "id": f"chatcmpl_{uuid.uuid4()}",
                "object": "chat.completion.chunk",
                "created": int(time.time()),
                "model": body.get("model", "openclaw"),
            }

            async for line in resp.aiter_lines():
                if not line:
                    continue
                if not line.startswith("data: "):
                    yield f"{line}\n".encode()
                    continue

                raw = line[6:]
                if raw == "[DONE]":
                    if not terminal_chunk_sent:
                        yield (
                            f"data: {json.dumps({**chunk_meta, 'choices': [{'index': 0, 'delta': {}, 'finish_reason': 'stop'}]})}\n\n"
                        ).encode()
                    yield b"data: [DONE]\n\n"
                    return

                try:
                    payload = json.loads(raw)
                except json.JSONDecodeError:
                    yield f"data: {raw}\n\n".encode()
                    continue

                for field in ("id", "object", "created", "model"):
                    if payload.get(field) is not None:
                        chunk_meta[field] = payload[field]

                choices = payload.get("choices") or []
                normalized_choices = []
                for idx, choice in enumerate(choices):
                    normalized_choice = dict(choice)
                    normalized_choice.setdefault("index", idx)
                    normalized_choice.setdefault("delta", {})
                    normalized_choice.setdefault("finish_reason", None)
                    if normalized_choice.get("finish_reason") is not None:
                        terminal_chunk_sent = True
                    normalized_choices.append(normalized_choice)

                normalized_payload = {
                    **chunk_meta,
                    **payload,
                    "choices": normalized_choices,
                }
                yield f"data: {json.dumps(normalized_payload)}\n\n".encode()

        async def stream():
            async def ollama_fallback_stream():
                logger.warning("OpenClaw provider chain exhausted; falling back to Ollama model %s", OLLAMA_MODEL)
                METRIC_LLM_REQUESTS.labels(backend="openclaw", status="fallback_ollama").inc()
                fallback_response = await call_ollama({**body, "model": OLLAMA_MODEL, "stream": True})
                async for chunk in fallback_response.body_iterator:
                    yield chunk

            try:
                METRIC_LLM_REQUESTS.labels(backend="openclaw", status="started").inc()
                async with httpx.AsyncClient(timeout=300) as c:
                    async with c.stream("POST", url, json=body, headers=headers) as resp:
                        if resp.status_code != 200:
                            err = await resp.aread()
                            logger.error("OpenClaw %s: %s", resp.status_code, err[:200])
                            if resp.status_code >= 500:
                                async for chunk in ollama_fallback_stream():
                                    yield chunk
                                return
                            METRIC_LLM_REQUESTS.labels(backend="openclaw", status="error").inc()
                            yield f'data: {{"error": "OpenClaw {resp.status_code}"}}\n\n'.encode()
                            return
                        METRIC_LLM_REQUESTS.labels(backend="openclaw", status="success").inc()
                        async for chunk in normalized_openclaw_stream(resp):
                            yield chunk
            except (httpx.TimeoutException, httpx.NetworkError) as exc:
                logger.error("OpenClaw stream error: %s", exc)
                async for chunk in ollama_fallback_stream():
                    yield chunk
            except Exception as exc:
                logger.error("OpenClaw stream error: %s", exc)
                METRIC_LLM_REQUESTS.labels(backend="openclaw", status="exception").inc()
                yield b'data: {"error": "proxy error"}\n\n'
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        try:
            async with httpx.AsyncClient(timeout=300) as c:
                resp = await c.post(url, json=body, headers=headers)
                resp.raise_for_status()
            METRIC_LLM_REQUESTS.labels(backend="openclaw", status="success").inc()
            METRIC_LLM_DURATION.labels(backend="openclaw").observe(time.monotonic() - t0)
            return JSONResponse(content=resp.json(), status_code=resp.status_code)
        except httpx.HTTPStatusError as exc:
            error_text = exc.response.text[:500]
            logger.error("OpenClaw %s: %s", exc.response.status_code, error_text)
            if exc.response.status_code >= 500:
                logger.warning("OpenClaw provider chain exhausted; falling back to Ollama model %s", OLLAMA_MODEL)
                METRIC_LLM_REQUESTS.labels(backend="openclaw", status="fallback_ollama").inc()
                return await call_ollama({**body, "model": OLLAMA_MODEL, "stream": False})
            METRIC_LLM_REQUESTS.labels(backend="openclaw", status="error").inc()
            raise
        except (httpx.TimeoutException, httpx.NetworkError) as exc:
            logger.error("OpenClaw transport error: %s", exc)
            logger.warning("OpenClaw unreachable; falling back to Ollama model %s", OLLAMA_MODEL)
            METRIC_LLM_REQUESTS.labels(backend="openclaw", status="fallback_ollama").inc()
            return await call_ollama({**body, "model": OLLAMA_MODEL, "stream": False})
        except Exception:
            METRIC_LLM_REQUESTS.labels(backend="openclaw", status="error").inc()
            raise


async def call_openai(body: dict):
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}", "Content-Type": "application/json"}
    if body.get("stream"):
        async def stream():
            for attempt in range(LLM_MAX_RETRIES + 1):
                try:
                    METRIC_LLM_REQUESTS.labels(backend="openai", status="started").inc()
                    async with httpx.AsyncClient(timeout=120) as c:
                        async with c.stream("POST", "https://api.openai.com/v1/chat/completions",
                                            json=body, headers=headers) as resp:
                            if resp.status_code != 200:
                                err = await resp.aread()
                                if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                                    ra = resp.headers.get("retry-after")
                                    wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                                    logger.warning("OpenAI %s (attempt %d/%d), retrying in %.1fs",
                                                   resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                                    await asyncio.sleep(wait)
                                    continue
                                logger.error("OpenAI %s: %s", resp.status_code, err[:200])
                                METRIC_LLM_REQUESTS.labels(backend="openai", status="error").inc()
                                yield f'data: {{"error": "OpenAI {resp.status_code}"}}\n\n'.encode()
                                return
                            METRIC_LLM_REQUESTS.labels(backend="openai", status="success").inc()
                            async for chunk in resp.aiter_bytes():
                                yield chunk
                            return
                except Exception as exc:
                    if attempt < LLM_MAX_RETRIES:
                        wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                        logger.warning("OpenAI stream error (attempt %d/%d): %s, retrying in %.1fs",
                                       attempt + 1, LLM_MAX_RETRIES + 1, exc, wait)
                        await asyncio.sleep(wait)
                        continue
                    logger.error("OpenAI stream error: %s", exc)
                    METRIC_LLM_REQUESTS.labels(backend="openai", status="exception").inc()
                    yield b'data: {"error": "proxy error"}\n\n'
                    return
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        for attempt in range(LLM_MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=120) as c:
                    resp = await c.post("https://api.openai.com/v1/chat/completions",
                                        json=body, headers=headers)
                    if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                        ra = resp.headers.get("retry-after")
                        wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                        logger.warning("OpenAI %s (attempt %d/%d), retrying in %.1fs",
                                       resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                        await asyncio.sleep(wait)
                        continue
                    resp.raise_for_status()
                METRIC_LLM_REQUESTS.labels(backend="openai", status="success").inc()
                METRIC_LLM_DURATION.labels(backend="openai").observe(time.monotonic() - t0)
                return JSONResponse(content=resp.json(), status_code=resp.status_code)
            except httpx.HTTPStatusError:
                METRIC_LLM_REQUESTS.labels(backend="openai", status="error").inc()
                raise
            except Exception:
                if attempt < LLM_MAX_RETRIES:
                    wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                    logger.warning("OpenAI non-stream error (attempt %d/%d), retrying in %.1fs",
                                   attempt + 1, LLM_MAX_RETRIES + 1, wait)
                    await asyncio.sleep(wait)
                    continue
                METRIC_LLM_REQUESTS.labels(backend="openai", status="error").inc()
                raise


def _is_retryable(status_code: int) -> bool:
    """Return True for status codes that warrant a retry (rate-limit, overloaded, server error)."""
    return status_code in (429, 529) or status_code >= 500


async def call_copilot(body: dict):
    """Forward chat completions to the GitHub Copilot OpenAI-compatible endpoint."""
    url = f"{COPILOT_BASE_URL}/chat/completions"
    headers = {
        "Authorization": f"Bearer {COPILOT_API_KEY}",
        "Content-Type": "application/json",
        "Copilot-Integration-Id": "gcor-proxy",
    }
    model = body.get("model", COPILOT_CHAT_MODEL)
    # Strip provider prefix if passed (e.g. "github-copilot/gpt-4.1" → "gpt-4.1")
    if "/" in model:
        model = model.split("/", 1)[1]
    body = {**body, "model": model}
    if body.get("stream"):
        async def stream():
            for attempt in range(LLM_MAX_RETRIES + 1):
                try:
                    METRIC_LLM_REQUESTS.labels(backend="copilot", status="started").inc()
                    async with httpx.AsyncClient(timeout=120) as c:
                        async with c.stream("POST", url, json=body, headers=headers) as resp:
                            if resp.status_code != 200:
                                err = await resp.aread()
                                if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                                    ra = resp.headers.get("retry-after")
                                    wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                                    logger.warning("Copilot %s (attempt %d/%d), retrying in %.1fs",
                                                   resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                                    await asyncio.sleep(wait)
                                    continue
                                logger.error("Copilot %s: %s", resp.status_code, err[:200])
                                METRIC_LLM_REQUESTS.labels(backend="copilot", status="error").inc()
                                yield f'data: {{"error": "Copilot {resp.status_code}"}}\n\n'.encode()
                                return
                            METRIC_LLM_REQUESTS.labels(backend="copilot", status="success").inc()
                            async for chunk in resp.aiter_bytes():
                                yield chunk
                            return
                except Exception as exc:
                    if attempt < LLM_MAX_RETRIES:
                        wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                        logger.warning("Copilot stream error (attempt %d/%d): %s, retrying in %.1fs",
                                       attempt + 1, LLM_MAX_RETRIES + 1, exc, wait)
                        await asyncio.sleep(wait)
                        continue
                    logger.error("Copilot stream error: %s", exc)
                    METRIC_LLM_REQUESTS.labels(backend="copilot", status="exception").inc()
                    yield b'data: {"error": "proxy error"}\n\n'
                    return
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        for attempt in range(LLM_MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=120) as c:
                    resp = await c.post(url, json=body, headers=headers)
                    if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                        ra = resp.headers.get("retry-after")
                        wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                        logger.warning("Copilot %s (attempt %d/%d), retrying in %.1fs",
                                       resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                        await asyncio.sleep(wait)
                        continue
                    resp.raise_for_status()
                METRIC_LLM_REQUESTS.labels(backend="copilot", status="success").inc()
                METRIC_LLM_DURATION.labels(backend="copilot").observe(time.monotonic() - t0)
                return JSONResponse(content=resp.json(), status_code=resp.status_code)
            except httpx.HTTPStatusError:
                METRIC_LLM_REQUESTS.labels(backend="copilot", status="error").inc()
                raise
            except Exception:
                if attempt < LLM_MAX_RETRIES:
                    wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                    logger.warning("Copilot non-stream error (attempt %d/%d), retrying in %.1fs",
                                   attempt + 1, LLM_MAX_RETRIES + 1, wait)
                    await asyncio.sleep(wait)
                    continue
                METRIC_LLM_REQUESTS.labels(backend="copilot", status="error").inc()
                raise


async def call_anthropic(body: dict):
    messages   = body.get("messages", [])
    system_txt = next((m["content"] for m in messages if m.get("role") == "system"), None)
    user_msgs  = [m for m in messages if m.get("role") != "system"]
    ant_body   = {"model": body.get("model", ANTHROPIC_CHAT_MODEL),
                  "max_tokens": body.get("max_tokens", 4096), "messages": user_msgs}
    if system_txt:
        ant_body["system"] = system_txt
    headers = {"x-api-key": ANTHROPIC_API_KEY, "anthropic-version": "2023-06-01",
               "Content-Type": "application/json"}

    if body.get("stream"):
        ant_body["stream"] = True
        async def stream():
            last_err = None
            for attempt in range(LLM_MAX_RETRIES + 1):
                try:
                    METRIC_LLM_REQUESTS.labels(backend="anthropic", status="started").inc()
                    async with httpx.AsyncClient(timeout=120) as c:
                        async with c.stream("POST", "https://api.anthropic.com/v1/messages",
                                            json=ant_body, headers=headers) as resp:
                            if resp.status_code != 200:
                                err = await resp.aread()
                                last_err = f"Anthropic {resp.status_code}"
                                if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                                    ra = resp.headers.get("retry-after")
                                    wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                                    logger.warning("Anthropic %s (attempt %d/%d), retrying in %.1fs",
                                                   resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                                    await asyncio.sleep(wait)
                                    continue
                                logger.error("Anthropic %s: %s", resp.status_code, err[:200])
                                METRIC_LLM_REQUESTS.labels(backend="anthropic", status="error").inc()
                                yield f'data: {{"error": "Anthropic {resp.status_code}"}}\n\n'.encode()
                                return
                            METRIC_LLM_REQUESTS.labels(backend="anthropic", status="success").inc()
                            async for line in resp.aiter_lines():
                                if not line.startswith("data: "):
                                    continue
                                raw = line[6:]
                                if raw == "[DONE]":
                                    yield b"data: [DONE]\n\n"; continue
                                try:
                                    evt = json.loads(raw)
                                    if evt.get("type") == "content_block_delta":
                                        txt = evt.get("delta", {}).get("text", "")
                                        yield f'data: {json.dumps({"choices": [{"delta": {"content": txt}, "finish_reason": None}]})}\n\n'.encode()
                                    elif evt.get("type") == "message_stop":
                                        yield b"data: [DONE]\n\n"
                                except json.JSONDecodeError:
                                    pass
                            return
                except Exception as exc:
                    last_err = str(exc)
                    if attempt < LLM_MAX_RETRIES:
                        wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                        logger.warning("Anthropic stream error (attempt %d/%d): %s, retrying in %.1fs",
                                       attempt + 1, LLM_MAX_RETRIES + 1, exc, wait)
                        await asyncio.sleep(wait)
                        continue
                    logger.error("Anthropic stream error: %s", exc)
                    METRIC_LLM_REQUESTS.labels(backend="anthropic", status="exception").inc()
                    yield b'data: {"error": "proxy error"}\n\n'
                    return
        return StreamingResponse(stream(), media_type="text/event-stream")
    else:
        t0 = time.monotonic()
        for attempt in range(LLM_MAX_RETRIES + 1):
            try:
                async with httpx.AsyncClient(timeout=120) as c:
                    resp = await c.post("https://api.anthropic.com/v1/messages",
                                        json=ant_body, headers=headers)
                    if _is_retryable(resp.status_code) and attempt < LLM_MAX_RETRIES:
                        ra = resp.headers.get("retry-after")
                        wait = float(ra) if ra else (LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5))
                        logger.warning("Anthropic %s (attempt %d/%d), retrying in %.1fs",
                                       resp.status_code, attempt + 1, LLM_MAX_RETRIES + 1, wait)
                        await asyncio.sleep(wait)
                        continue
                    resp.raise_for_status()
                METRIC_LLM_REQUESTS.labels(backend="anthropic", status="success").inc()
                METRIC_LLM_DURATION.labels(backend="anthropic").observe(time.monotonic() - t0)
                ant = resp.json()
                usage = ant.get("usage", {})
                return JSONResponse(content={
                    "id": ant.get("id", ""), "object": "chat.completion",
                    "model": ant.get("model", ant_body["model"]),
                    "choices": [{"index": 0,
                                 "message": {"role": "assistant",
                                             "content": ant.get("content", [{}])[0].get("text", "")},
                                 "finish_reason": ant.get("stop_reason", "stop")}],
                    "usage": {"prompt_tokens":    usage.get("input_tokens", 0),
                              "completion_tokens": usage.get("output_tokens", 0),
                              "total_tokens":      usage.get("input_tokens", 0) + usage.get("output_tokens", 0)},
                })
            except httpx.HTTPStatusError:
                METRIC_LLM_REQUESTS.labels(backend="anthropic", status="error").inc()
                raise
            except Exception:
                if attempt < LLM_MAX_RETRIES:
                    wait = LLM_RETRY_BASE * (2 ** attempt) + random.uniform(0, 0.5)
                    logger.warning("Anthropic non-stream error (attempt %d/%d), retrying in %.1fs",
                                   attempt + 1, LLM_MAX_RETRIES + 1, wait)
                    await asyncio.sleep(wait)
                    continue
                METRIC_LLM_REQUESTS.labels(backend="anthropic", status="error").inc()
                raise


# ── Routes ─────────────────────────────────────────────────────────────────────

@app.get("/v1/models")
async def list_models():
    models = []
    # Pull Ollama models
    if LLM_BACKEND == "ollama" or EMBEDDING_BACKEND == "ollama":
        try:
            async with httpx.AsyncClient(timeout=5) as c:
                resp = await c.get(f"{OLLAMA_BASE_URL}/api/tags")
                if resp.status_code == 200:
                    for m in resp.json().get("models", []):
                        models.append({"id": m["name"], "object": "model", "owned_by": "ollama"})
        except Exception as exc:
            logger.warning("Could not fetch Ollama models: %s", exc)
    # Pull OpenAI models
    if OPENAI_API_KEY:
        try:
            async with httpx.AsyncClient(timeout=10) as c:
                resp = await c.get("https://api.openai.com/v1/models",
                                   headers={"Authorization": f"Bearer {OPENAI_API_KEY}"})
                if resp.status_code == 200:
                    models += resp.json().get("data", [])
        except Exception as exc:
            logger.warning("Could not fetch OpenAI models: %s", exc)
    if not models:
        models = [
            {"id": OPENAI_CHAT_MODEL,    "object": "model", "owned_by": "openai"},
            {"id": ANTHROPIC_CHAT_MODEL, "object": "model", "owned_by": "anthropic"},
            {"id": OLLAMA_MODEL,         "object": "model", "owned_by": "ollama"},
        ]
    models = [{"id": "openclaw", "object": "model", "created": 1700000000, "owned_by": "eedgeai"}] + models
    return {"object": "list", "data": models}


@app.post("/v1/embeddings")
async def embeddings(request: Request):
    """Proxy embedding requests to the configured backend (for OpenWebUI's built-in RAG)."""
    body = await request.json()
    try:
        if EMBEDDING_BACKEND == "ollama":
            async with httpx.AsyncClient(timeout=60) as c:
                resp = await c.post(f"{OLLAMA_BASE_URL}/v1/embeddings", json=body)
                resp.raise_for_status()
                return JSONResponse(content=resp.json())
        if not OPENAI_API_KEY:
            raise HTTPException(status_code=503, detail="OPENAI_API_KEY not configured")
        async with httpx.AsyncClient(timeout=30) as c:
            resp = await c.post(
                "https://api.openai.com/v1/embeddings",
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}",
                         "Content-Type": "application/json"},
                json=body,
            )
            resp.raise_for_status()
            return JSONResponse(content=resp.json())
    except httpx.HTTPStatusError as exc:
        raise HTTPException(status_code=exc.response.status_code,
                            detail="Embeddings upstream error")
    except Exception as exc:
        logger.error("Embeddings failed: %s", exc)
        raise HTTPException(status_code=502, detail="Embeddings request failed")


@app.post("/v1/chat/completions")
async def chat_completions(request: Request):
    body = await request.json()
    rag_collection = _request_rag_collection(body, request)
    return await _run_chat_completion(body, rag_collection)


@app.post("/v1/openclaw/chat/completions")
async def openclaw_chat_completions_raw(request: Request):
    """Raw passthrough to openclaw, with its existing Ollama-on-failure
    fallback (call_openclaw already falls back on 5xx/timeout/network
    error) — but WITHOUT /v1/chat/completions' GCOR context-injection
    pipeline (Graphiti search + prompt augmentation), which would corrupt
    a caller's own carefully-built prompt (e.g. graphiti-core's own
    entity/edge extraction schema instructions). For backend workloads
    that want openclaw's speed with Ollama resilience, not the GCOR
    RAG behavior meant for end-user chat turns. See
    docs/adr/0010-graphiti-openclaw-llm.md.
    """
    body = await request.json()
    return await call_openclaw(body)


# ── Buzz bridge ──────────────────────────────────────────────────────────────
# In-process sliding-window limiter — the Buzz bridge is a single trusted
# caller, this just bounds cost/blast-radius if its token ever leaks or the
# bridge misbehaves (e.g. a workflow loop re-triggering itself).
_buzz_rate_lock: "asyncio.Lock | None" = None
_buzz_rate_window: list[float] = []


async def _check_buzz_rate_limit():
    global _buzz_rate_lock
    if _buzz_rate_lock is None:
        _buzz_rate_lock = asyncio.Lock()
    now = time.monotonic()
    async with _buzz_rate_lock:
        while _buzz_rate_window and now - _buzz_rate_window[0] > 60:
            _buzz_rate_window.pop(0)
        if len(_buzz_rate_window) >= BUZZ_BRIDGE_RATE_LIMIT:
            raise HTTPException(status_code=429, detail="Buzz bridge rate limit exceeded")
        _buzz_rate_window.append(now)


@app.post("/v1/buzz/chat/completions")
async def buzz_chat_completions(request: Request):
    """Bearer-authenticated GCOR entry point for the Buzz collaboration bridge.

    Buzz channel content (human or agent authored) is a less-trusted source than
    the rest of the internal docker network, so this route is separate from
    /v1/chat/completions: it requires BUZZ_BRIDGE_API_KEY, caps message size,
    optionally restricts which collection may be queried, and rate-limits.
    """
    if not BUZZ_BRIDGE_API_KEY:
        raise HTTPException(status_code=503, detail="Buzz bridge is not configured")

    auth = request.headers.get("authorization", "")
    scheme, _, token = auth.partition(" ")
    if scheme.lower() != "bearer" or not hmac.compare_digest(token, BUZZ_BRIDGE_API_KEY):
        raise HTTPException(status_code=401, detail="Unauthorized")

    await _check_buzz_rate_limit()

    body = await request.json()
    messages = body.get("messages", [])
    if not isinstance(messages, list) or not messages or len(messages) > 50:
        raise HTTPException(status_code=400, detail="Invalid messages")
    for m in messages:
        content = m.get("content", "") if isinstance(m, dict) else None
        # Same two shapes last_user_text() already tolerates: a plain string, or
        # OpenAI-style content-part blocks ([{"type": "text", "text": "..."}]) —
        # buzz-agent's OpenAI-compat client sends the latter.
        if isinstance(content, str):
            text_len = len(content)
        elif isinstance(content, list):
            text_len = sum(
                len(p.get("text", "")) for p in content
                if isinstance(p, dict) and p.get("type") == "text"
            )
        else:
            raise HTTPException(status_code=400, detail="Invalid message content")
        if text_len > BUZZ_BRIDGE_MAX_MESSAGE_CHARS:
            raise HTTPException(status_code=400, detail="Message content too long")

    rag_collection = _request_rag_collection(body, request)
    if BUZZ_ALLOWED_COLLECTIONS and rag_collection not in (None, *BUZZ_ALLOWED_COLLECTIONS):
        raise HTTPException(status_code=403, detail="Collection not permitted for Buzz bridge")

    body = {**body, "model": "openclaw"}  # always answer via the GCOR pipeline, ignore caller-supplied model
    return await _run_chat_completion(body, rag_collection)


@app.post("/v1/buzz/ingest")
async def buzz_ingest(
    request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile | None = File(None),
    storage_key: str = Form(""),
):
    """Buzz-facing document ingestion — same trust boundary and auth as
    /v1/buzz/chat/completions (BUZZ_BRIDGE_API_KEY, shared rate limit).

    Two ways in, both used by the bridge:
      - multipart file upload — an image attachment pulled off a Buzz message
        (Buzz's own media pipeline only accepts images, so this is the only
        binary content a channel can actually produce).
      - storage_key — ingest (or re-confirm) a file that's already sitting in
        MinIO under documents/inbox/, processed/, failed/, or originals/*/,
        for the '@bot ingest <name>' text-reference path.
    """
    if not BUZZ_BRIDGE_API_KEY:
        raise HTTPException(status_code=503, detail="Buzz bridge is not configured")

    auth = request.headers.get("authorization", "")
    scheme, _, token = auth.partition(" ")
    if scheme.lower() != "bearer" or not hmac.compare_digest(token, BUZZ_BRIDGE_API_KEY):
        raise HTTPException(status_code=401, detail="Unauthorized")

    await _check_buzz_rate_limit()

    if file is not None:
        data = await file.read()
        filename = file.filename or "upload"
    elif storage_key.strip():
        data, filename = await _find_and_fetch_from_s3(storage_key.strip())
        if data is None:
            raise HTTPException(
                status_code=404,
                detail=f"'{storage_key}' not found in MinIO (checked bucket root, inbox/, processed/, failed/, originals/*)",
            )
    else:
        raise HTTPException(status_code=400, detail="Provide either a file or storage_key")

    if not data:
        raise HTTPException(status_code=400, detail="Empty file")
    if len(data) > BUZZ_BRIDGE_MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=400, detail="File too large")

    return await _ingest_bytes(filename, data, background_tasks=background_tasks, enable_docint=True)


async def _run_chat_completion(body: dict, rag_collection: str | None):
    messages = list(body.get("messages", []))

    rag_intent   = "none"
    rag_fallback = "disabled"
    rag_t0       = time.monotonic()

    if ENABLE_RAG and messages:
        query = last_user_text(messages)

        if query:
            # ── Step 1: Intent ──────────────────────────────────────────────
            intent     = await classify_intent(query)
            rag_intent = intent
            logger.info("GCOR intent: %s", intent)

            # ── Step 2: Graphiti hybrid temporal retrieval ──────────────────
            raw_hits, graph_records = await graphiti_search(query, rag_collection)
            qdrant_hits = _filter_hits(raw_hits)
            METRIC_GRAPHITI_FACTS.observe(len(qdrant_hits))
            logger.info(
                "Graphiti[%s]: %d facts → %d after policy filters",
                rag_collection or GRAPHITI_GROUP_ID, len(raw_hits), len(qdrant_hits),
            )
            METRIC_CONTEXT_RECORDS.observe(len(graph_records))
            logger.info("Graphiti: %d temporal graph records retrieved", len(graph_records))

            # ── Step 2c: Belief conflict resolution ─────────────────────────
            graph_records = resolve_belief_conflicts(graph_records)

            # Graphiti fact results retain their source episode UUIDs.
            provenance = " | ".join(
                str((hit.get("payload") or {}).get("document_id", ""))
                for hit in qdrant_hits[:3]
                if (hit.get("payload") or {}).get("document_id")
            )
            if provenance:
                logger.info("Provenance chain: %s", provenance[:120])

            # ── Step 3: Reflection + context build ──────────────────────────
            context  = build_gcor_context(intent, qdrant_hits, graph_records, provenance=provenance)

            # Determine fallback mode for metrics
            if graph_records:
                rag_fallback = "graphiti"
            elif qdrant_hits:
                rag_fallback = "graphiti_fact"
            else:
                rag_fallback = "llm_only"

            messages = inject_context(messages, context)
            body     = {**body, "messages": messages}

        METRIC_RAG_REQUESTS.labels(intent=rag_intent, fallback_mode=rag_fallback).inc()
        METRIC_RAG_DURATION.observe(time.monotonic() - rag_t0)

    # ── Route to LLM ──────────────────────────────────────────────────────────
    model, backend = _resolve_model(body.get("model", "openclaw"))
    body = {**body, "model": model}
    try:
        if backend == "openclaw":
            return await call_openclaw(body)
        elif backend == "ollama":
            return await call_ollama(body)
        elif backend == "anthropic" and ANTHROPIC_API_KEY:
            return await call_anthropic(body)
        elif OPENAI_API_KEY:
            return await call_openai(body)
        else:
            raise HTTPException(status_code=503, detail="No LLM backend configured")
    except httpx.HTTPStatusError as exc:
        logger.error("LLM upstream error: %s", exc)
        raise HTTPException(status_code=exc.response.status_code, detail="LLM upstream error")
    except HTTPException:
        raise
    except Exception as exc:
        logger.error("Chat completion failed: %s", exc)
        raise HTTPException(status_code=502, detail="Chat completion failed")


_TEMPLATES = os.path.join(os.path.dirname(__file__), "templates")


# ── Knowledge UI ───────────────────────────────────────────────────────────────

@app.get("/", response_class=RedirectResponse, status_code=302)
async def root():
    return "/knowledge"


@app.get("/openclaw", response_class=RedirectResponse, status_code=302)
async def openclaw_redirect():
    token = OPENCLAW_GATEWAY_TOKEN
    base  = "http://127.0.0.1:18789"
    if token:
        return f"{base}/#token={token}"
    return base


@app.get("/knowledge", response_class=HTMLResponse)
async def knowledge_ui():
    with open(os.path.join(_TEMPLATES, "knowledge.html"), encoding="utf-8") as f:
        return f.read()


@app.get("/api/buckets")
async def api_buckets():
    """List MinIO buckets, so the ingest UI can offer a picker alongside a
    free-text name to create a new one (the ingest endpoint auto-creates
    whatever bucket name it's given)."""
    client = _get_s3_client()
    if client is None:
        return {"buckets": [], "default": S3_BUCKET}
    try:
        resp = await asyncio.to_thread(client.list_buckets)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"MinIO: {exc}")
    buckets = [
        {"name": b["Name"], "created_at": b["CreationDate"].isoformat()}
        for b in resp.get("Buckets", [])
    ]
    if not any(b["name"] == S3_BUCKET for b in buckets):
        buckets.insert(0, {"name": S3_BUCKET, "created_at": None})
    return {"buckets": buckets, "default": S3_BUCKET}


@app.get("/api/collections")
async def api_collections():
    """List configured Graphiti group namespaces and recent episodes."""
    names = list(dict.fromkeys(
        [GRAPHITI_GROUP_ID, *RAG_EXTRA_COLLECTIONS, GRAPHITI_CHAT_SESSIONS_GROUP_ID]
    ))
    collections = []
    async with httpx.AsyncClient(timeout=30) as client:
        for name in names:
            recent = []
            count = 0
            try:
                response = await client.get(f"{GRAPHITI_URL}/episodes/{name}", params={"last_n": 200})
                response.raise_for_status()
                episodes = response.json()
                count = len(episodes)
                recent = [
                    {
                        "doc_id": episode.get("uuid", ""),
                        "title": episode.get("name") or "Graphiti episode",
                        "created_at": episode.get("created_at") or episode.get("valid_at"),
                        "access_level": "public",
                    }
                    for episode in episodes[:5]
                ]
            except Exception as exc:
                logger.warning("Graphiti group metadata failed for '%s': %s", name, exc)
            collections.append({
                "name": name,
                "points_count": count,
                "doc_count": count,
                "recent_docs": recent,
                "embedding_model": os.getenv("GRAPHITI_EMBEDDING_MODEL", OLLAMA_EMBEDDING_MODEL),
            })
    return {"collections": collections, "embedding_model": os.getenv("GRAPHITI_EMBEDDING_MODEL", OLLAMA_EMBEDDING_MODEL)}


@app.get("/api/collections/{name}/docs")
async def api_collection_docs(name: str):
    """List recent Graphiti episodes for a group namespace."""
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            response = await client.get(f"{GRAPHITI_URL}/episodes/{name}", params={"last_n": 200})
            response.raise_for_status()
            episodes = response.json()
        docs = [
            {
                "doc_id": episode.get("uuid", ""),
                "title": episode.get("name") or "Graphiti episode",
                "created_at": episode.get("created_at") or episode.get("valid_at"),
                "access_level": "public",
                "chunk_count": 1,
            }
            for episode in episodes
        ]
        return {"collection": name, "docs": docs}
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc))


@app.delete("/api/collections/{name}/docs/{doc_id}")
async def api_archive_document(name: str, doc_id: str):
    """Delete one Graphiti episode. Source files remain preserved in MinIO.

    This is a permanent deletion from the knowledge graph — Graphiti has no
    archive/restore concept (see docs/adr/0001-graphiti-falkordb-backend.md).
    To bring a document back, re-ingest its original from MinIO.
    """
    async with httpx.AsyncClient(timeout=30) as client:
        response = await client.delete(f"{GRAPHITI_URL}/episode/{doc_id}")
        if response.status_code == 404:
            raise HTTPException(status_code=404, detail=f"Episode '{doc_id}' not found")
        response.raise_for_status()
    return {"status": "deleted", "doc_id": doc_id, "collection": name}


@app.post("/api/collections")
async def api_create_collection(request: Request):
    """Validate a new Graphiti group name; groups materialize on first ingest."""
    body = await request.json()
    name = (body.get("name") or "").strip()
    if not name or "/" in name or ".." in name:
        raise HTTPException(status_code=400, detail="Invalid collection name")
    METRIC_COLLECTION_OPS.labels(operation="create").inc()
    return {"status": "ok", "name": name, "materializes_on_first_ingest": True}


@app.patch("/api/collections/{name}")
async def api_rename_collection(name: str, request: Request):
    """Graphiti group renames require re-ingestion to preserve provenance."""
    body = await request.json()
    new_name = (body.get("name") or "").strip()
    if not new_name or "/" in new_name or ".." in new_name:
        raise HTTPException(status_code=400, detail="Invalid collection name")
    raise HTTPException(
        status_code=409,
        detail="Graphiti group IDs are immutable; ingest into the new group, verify it, then delete the old group.",
    )


@app.delete("/api/collections/{name}")
async def api_archive_collection(name: str):
    """Delete a Graphiti group after its source objects have been retained in MinIO.

    This is a permanent deletion — Graphiti has no archive/restore concept
    (see docs/adr/0001-graphiti-falkordb-backend.md). To bring a group back,
    re-ingest its source documents from MinIO into a new group.
    """
    if name == GRAPHITI_GROUP_ID:
        raise HTTPException(status_code=409, detail="The default Graphiti group cannot be deleted")
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.delete(f"{GRAPHITI_URL}/group/{name}")
        response.raise_for_status()
    METRIC_COLLECTION_OPS.labels(operation="delete").inc()
    return {"status": "deleted", "name": name}


@app.get("/api/search")
async def api_search(
    collection: str = Query(...),
    q: str = Query(...),
    top_k: int = Query(8),
):
    """Hybrid temporal search against a Graphiti group namespace."""
    try:
        hits, _ = await graphiti_search(q, collection)
    except Exception as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    # Apply the same confidence/temporal-validity/access-control filtering the
    # chat RAG path uses (_run_chat_completion) — this endpoint used to return
    # raw hits, which meant a document marked access_level="restricted" (or
    # expired, or below CONFIDENCE_THRESHOLD) was hidden from chat but still
    # fully readable here. Keep both paths honoring the same policy.
    hits = _filter_hits(hits)

    results = [
        {
            "score":            h.get("score", 0),
            "text":             h.get("payload", {}).get("text", ""),
            "graphiti_edge_uuid": h.get("payload", {}).get("graphiti_edge_uuid", ""),
            "doc_title":        h.get("payload", {}).get("document_title", ""),
            "position":         h.get("payload", {}).get("position"),
        }
        for h in hits[:top_k]
    ]
    return {"collection": collection, "query": q, "results": results}


async def _ingest_bytes(
    filename: str,
    data: bytes,
    *,
    background_tasks: BackgroundTasks | None = None,
    title: str = "",
    agent_id: str = "",
    access_level: str = "public",
    enable_docint: bool = False,
    collection: str = "",
    bucket: str = "",
    valid_hours: float = 0.0,
) -> dict:
    """Core ingest pipeline shared by every entry point: extract -> chunk ->
    Graphiti (FalkorDB), plus a best-effort original-file copy to MinIO.
    Raises HTTPException on failure."""
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")
    try:
        access_level = validate_access_level(access_level)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))

    ext = _resolve_ext(filename, data)
    if ext not in _INGEST_ACCEPT:
        raise HTTPException(status_code=415, detail=f"Unsupported type '{ext}'")

    doc_title  = title.strip() or filename
    doc_id     = hashlib.md5(f"{filename}-{datetime.now().isoformat()}".encode()).hexdigest()[:16]
    image_props: dict = {}
    docint_summary: dict = {}

    target_bucket = _validate_bucket_name(bucket.strip() or S3_BUCKET)
    storage_key = await _store_original_to_s3(doc_id, filename, data, bucket=target_bucket)

    async def _run_docint() -> str:
        nonlocal image_props, docint_summary
        result: DocIntelResult = await process_document(filename, data)
        image_props = result.image_props
        docint_summary = {
            "document_type":    result.document_type,
            "document_subtype": result.document_subtype,
            "is_scanned":       result.is_scanned,
            "pages":            result.pages,
            "language":         result.language,
            "tables":           len(result.tables),
            "form_fields":      len(result.form_fields),
            "entities": {
                "dates":         result.entities.dates[:5],
                "amounts":       result.entities.amounts[:5],
                "names":         result.entities.names[:5],
                "organizations": result.entities.organizations[:5],
            },
        }
        return result.to_rich_text().strip()

    # PDFs/DOCX where OCR was never asked for, but the plain-text layer turns
    # out empty (e.g. a scanned page with no embedded text — a photographed
    # ID or a screenshotted webpage saved as PDF), auto-retry through the
    # Document Intelligence/OCR pipeline before giving up. Before this,
    # `enable_docint` defaulted to false for every entry point including
    # `gcor_file_ingest` (no chat user can pass that flag), so any scanned
    # attachment failed with a silent 422 and the assistant reporting "I
    # don't see anything attached" — see
    # docs/adr/0009-ingest-ocr-auto-fallback.md. Text-native formats
    # (.txt/.md/.json/.csv) are excluded: empty output there means an
    # actually-empty file, not something OCR could rescue, so they still
    # fail fast instead of paying for a ~1-2 minute OCR pass for nothing.
    _OCR_FALLBACK_EXTS = {".pdf", ".docx"}

    try:
        if ext in _MEDICAL_EXTS or _is_image(ext):
            # Medical and image formats always use their dedicated handlers —
            # DocInt pipeline does not understand DICOM/NIfTI/pixel data.
            text, image_props = await _extract_image_text(filename, data, ext)
            text = text.strip()
        elif enable_docint:
            # Full Document Intelligence pipeline (PDFs, DOCX, scanned docs)
            text = await _run_docint()
        else:
            text = _extract_text(filename, data).strip()
            if not text and ext in _OCR_FALLBACK_EXTS:
                text = await _run_docint()
                enable_docint = True  # reflect what actually ran in the response/logs
    except Exception as exc:
        raise HTTPException(status_code=422, detail=f"Extraction failed: {exc}")

    if not text:
        raise HTTPException(status_code=422, detail="No text extracted")
    chunks = _chunk_text(text)
    if not chunks:
        raise HTTPException(status_code=422, detail="No chunks produced")
    target_collection = collection.strip() or QDRANT_COLLECTION

    if storage_key:
        image_props["storage_bucket"] = target_bucket
        image_props["storage_key"] = storage_key

    from datetime import timedelta
    valid_to = (
        (datetime.now(timezone.utc) + timedelta(hours=valid_hours)).isoformat()
        if valid_hours > 0 else None
    )
    logger.info(
        "Ingest '%s': %d chunks (ext=%s, docint=%s, collection=%s, bucket=%s, valid_to=%s)",
        doc_title, len(chunks), ext, enable_docint, target_collection, target_bucket, valid_to,
    )

    try:
        upserted = await graphiti_ingest(
            target_collection, doc_id, doc_title, filename, chunks, agent_id, access_level,
            valid_to=valid_to,
        )
    except Exception as exc:
        METRIC_INGEST_TOTAL.labels(status="graphiti_error").inc()
        raise HTTPException(status_code=502, detail=f"Graphiti: {exc}")

    # Graphiti extracts entities, relationships, provenance, and temporal facts
    # asynchronously from each episode; no separate NER/graph backfill is needed.

    METRIC_INGEST_TOTAL.labels(status="success").inc()
    METRIC_INGEST_CHUNKS.observe(len(chunks))

    result_json = {"status": "ok", "document_id": doc_id, "title": doc_title,
                   "filename": filename, "chunks": len(chunks), "graphiti_episodes": upserted,
                   "collection": target_collection}
    if valid_to:
        result_json["valid_to"] = valid_to
    if image_props:
        result_json["image_metadata"] = image_props
    if docint_summary:
        result_json["docint"] = docint_summary
    if storage_key:
        result_json["storage"] = {"bucket": target_bucket, "key": storage_key}
    return result_json


@app.post("/api/ingest")
async def api_ingest(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: str = Form(""),
    agent_id: str = Form(""),
    access_level: str = Form("public"),
    enable_docint: str = Form("false"),   # "true" → run full Document Intelligence
    collection: str = Form(""),           # target Qdrant collection (defaults to QDRANT_COLLECTION)
    bucket: str = Form(""),               # target MinIO bucket (defaults to S3_BUCKET, created if missing)
    valid_hours: float = Form(0.0),       # >0 → chunk expires after this many hours
):
    """JSON-returning ingest endpoint used by the knowledge UI."""
    filename = file.filename or "upload"
    data     = await file.read()
    return await _ingest_bytes(
        filename, data,
        background_tasks=background_tasks,
        title=title, agent_id=agent_id, access_level=access_level,
        enable_docint=enable_docint.lower() in ("true", "1", "yes"),
        collection=collection, bucket=bucket, valid_hours=valid_hours,
    )


# ── Chat session ingest ─────────────────────────────────────────────────────────
# Archives full chat-session transcripts to MinIO (S3) and indexes them into
# Graphiti, so conversation history is retrievable the same way documents are —
# not just files uploaded during a session. Called by the
# gcor_chat_session_ingest OpenWebUI filter after each assistant turn; see
# openwebui-functions/gcor_chat_session_ingest.py.

GRAPHITI_CHAT_SESSIONS_GROUP_ID = os.getenv("GRAPHITI_CHAT_SESSIONS_GROUP_ID", "chat_sessions")
S3_CHAT_SESSIONS_PREFIX         = os.getenv("S3_CHAT_SESSIONS_PREFIX", "chat-sessions/")

_SESSION_SLUG_RE = re.compile(r"[^a-z0-9]+")

# OpenWebUI stamps every new chat with one of these placeholder titles before
# its own async title-generation call replaces them a turn or two later.
# Treating them as "no real title yet" avoids archiving many unrelated
# sessions under the same generic name (every fresh chat landing on
# "new-chat"), and avoids a Graphiti episode literally titled "New Chat".
_GENERIC_SESSION_TITLES = {"", "new chat", "untitled", "untitled chat", "new conversation"}


def _slugify(text: str) -> str:
    """Sanitize free text into an S3-key-safe, still-human-readable slug."""
    return _SESSION_SLUG_RE.sub("-", (text or "").strip().lower()).strip("-")


def _meaningful_session_name(title: str, messages: list, fallback: str) -> str:
    """Pick a human-readable name for a chat session: the real title if
    OpenWebUI has generated one yet, else a snippet of the first user
    message, else the fallback (normally session_id)."""
    title = (title or "").strip()
    if title.lower() not in _GENERIC_SESSION_TITLES:
        return title
    for m in messages:
        if not (isinstance(m, dict) and m.get("role") == "user"):
            continue
        content = m.get("content")
        if isinstance(content, list):
            content = " ".join(
                part.get("text", "") for part in content
                if isinstance(part, dict) and part.get("type") == "text"
            )
        content = str(content or "").strip()
        if content:
            return content[:60]
    return fallback or "session"


def _render_transcript(messages: list) -> str:
    """Render chat messages as plain 'role: content' text for chunking/embedding."""
    lines = []
    for m in messages:
        if not isinstance(m, dict):
            continue
        role = str(m.get("role", "user")).strip() or "user"
        content = m.get("content")
        if isinstance(content, list):
            # OpenAI/OpenWebUI multi-part content (text + image parts, etc.)
            content = " ".join(
                part.get("text", "") for part in content
                if isinstance(part, dict) and part.get("type") == "text"
            )
        content = str(content or "").strip()
        if content:
            lines.append(f"{role}: {content}")
    return "\n\n".join(lines)


async def _store_chat_session_to_s3(
    key: str, payload: dict, bucket: str = S3_BUCKET,
) -> str | None:
    """Persist the full chat-session transcript (JSON) to MinIO at the given key.

    Best-effort like _store_original_to_s3 — returns None (never raises) if
    MinIO is unconfigured or unreachable.
    """
    client = _get_s3_client()
    if client is None:
        return None
    body = json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")
    try:
        await _ensure_bucket_exists(client, bucket)
        await asyncio.to_thread(
            client.put_object,
            Bucket=bucket, Key=key, Body=body, ContentType="application/json",
        )
        return key
    except Exception as exc:
        logger.warning("Chat session S3 store failed for key '%s': %s", key, exc)
        return None


@app.post("/api/ingest/session")
async def api_ingest_session(request: Request):
    """Archive one chat-session snapshot to MinIO and Graphiti.

    Expects JSON: {session_id, title, messages: [{role, content}, ...],
    model, agent_id, access_level, collection}. Every call writes a new
    snapshot of the full transcript — the OpenWebUI filter calls this after
    each assistant turn.

    Naming: the session's snapshots all live under one MinIO folder keyed by
    session_id — stable for the life of the conversation, even though
    OpenWebUI's own title starts as a generic placeholder ("New Chat") and
    only gets renamed by its own async title-generation a turn or two later.
    Each snapshot's filename carries the date, time, and the best
    human-readable name available at that moment (the real title once
    OpenWebUI has generated one, otherwise a snippet of the first user
    message) — e.g.
    chat-sessions/<session_id>/2026-08-22T14-05-30Z_what-is-the-capital-of-france.json
    """
    body = await request.json()
    session_id = str(body.get("session_id") or "").strip()
    title      = str(body.get("title") or "").strip()
    messages   = body.get("messages") or []
    model        = str(body.get("model") or "").strip()
    agent_id     = str(body.get("agent_id") or "").strip()
    try:
        access_level = validate_access_level(str(body.get("access_level") or "").strip())
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc))
    collection   = str(body.get("collection") or "").strip() or GRAPHITI_CHAT_SESSIONS_GROUP_ID

    if not isinstance(messages, list) or not messages:
        raise HTTPException(status_code=400, detail="messages must be a non-empty list")

    transcript_text = _render_transcript(messages)
    if not transcript_text:
        raise HTTPException(status_code=422, detail="No text content in messages")

    now       = datetime.now(timezone.utc)
    timestamp = now.isoformat()
    file_stamp = now.strftime("%Y-%m-%dT%H-%M-%SZ")

    session_name    = _meaningful_session_name(title, messages, session_id)
    name_slug       = _slugify(session_name)[:80] or "session"
    session_folder  = _slugify(session_id)[:80] or name_slug
    doc_title       = session_name

    intended_key = f"{S3_CHAT_SESSIONS_PREFIX}{session_folder}/{file_stamp}_{name_slug}.json"
    storage_key = await _store_chat_session_to_s3(
        intended_key,
        {
            "session_id":    session_id,
            "title":         title,
            "session_name":  session_name,
            "model":         model,
            "message_count": len(messages),
            "ingested_at":   timestamp,
            "messages":      messages,
        },
    )

    chunks = _chunk_text(transcript_text)
    if not chunks:
        raise HTTPException(status_code=422, detail="Transcript produced no chunks")
    doc_id = hashlib.md5(f"{session_folder}-{timestamp}".encode()).hexdigest()[:16]

    try:
        upserted = await graphiti_ingest(
            collection, doc_id, doc_title,
            f"chat_session:{session_id or session_folder}", chunks, agent_id, access_level,
        )
    except Exception as exc:
        METRIC_INGEST_TOTAL.labels(status="graphiti_error").inc()
        raise HTTPException(status_code=502, detail=f"Graphiti: {exc}")

    METRIC_INGEST_TOTAL.labels(status="success").inc()
    METRIC_INGEST_CHUNKS.observe(len(chunks))

    result = {
        "status":            "ok",
        "document_id":       doc_id,
        "title":             doc_title,
        "session_id":        session_id,
        "messages":          len(messages),
        "chunks":            len(chunks),
        "graphiti_episodes": upserted,
        "collection":        collection,
        "timestamp":         timestamp,
    }
    if storage_key:
        result["storage"] = {"bucket": S3_BUCKET, "key": storage_key}
    return result


# The one-time Neo4j/Qdrant migration and backfill endpoints that used to
# live here (ner-backfill, graph-backfill, graph-migrate-chunk-edges,
# buddy-memory-migrate) were removed with the Graphiti/FalkorDB cutover —
# see docs/adr/0001-graphiti-falkordb-backend.md. Graphiti extracts entities,
# relationships, provenance, and temporal facts automatically from each
# episode, so no separate backfill step exists in the new architecture.


@app.post("/api/docint")
async def api_docint(
    file: UploadFile = File(...),
    extract_tables:   str = Form("true"),
    extract_forms:    str = Form("true"),
    extract_entities: str = Form("true"),
    classify:         str = Form("true"),
    vision:           str = Form("true"),
):
    """
    Standalone Document Intelligence endpoint.
    Returns rich structured output without ingesting into Graphiti.
    """
    filename = file.filename or "upload"
    data     = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file")

    ext = _resolve_ext(filename, data)
    if ext not in _INGEST_ACCEPT:
        raise HTTPException(status_code=415, detail=f"Unsupported type '{ext}'")

    def _bool(v: str) -> bool:
        return v.lower() in ("true", "1", "yes")

    try:
        result: DocIntelResult = await process_document(
            filename, data,
            extract_tables=_bool(extract_tables),
            extract_forms=_bool(extract_forms),
            extract_entities=_bool(extract_entities),
            classify=_bool(classify),
            vision=_bool(vision),
        )
    except Exception as exc:
        raise HTTPException(status_code=422, detail=str(exc))

    return {
        "filename":         filename,
        "document_type":    result.document_type,
        "document_subtype": result.document_subtype,
        "type_confidence":  result.type_confidence,
        "language":         result.language,
        "is_scanned":       result.is_scanned,
        "pages":            result.pages,
        "char_count":       result.char_count,
        "tables": [
            {"caption": t.caption, "markdown": t.markdown,
             "rows": t.row_count, "cols": t.col_count, "page": t.page}
            for t in result.tables
        ],
        "form_fields": result.form_fields,
        "entities": {
            "dates":         result.entities.dates,
            "amounts":       result.entities.amounts,
            "names":         result.entities.names,
            "organizations": result.entities.organizations,
            "identifiers":   result.entities.identifiers,
            "locations":     result.entities.locations,
        },
        "text_preview": result.text[:500],
    }


@app.get("/health")
async def health():
    return {
        "status":               "ok",
        "rag":                  ENABLE_RAG,
        "confidence_threshold": CONFIDENCE_THRESHOLD,
        "agent_id":             AGENT_ID or None,
    }


# ── Document Ingest ────────────────────────────────────────────────────────────

_INGEST_ACCEPT = {
    # Documents
    ".txt", ".md", ".pdf", ".docx", ".json", ".csv",
    # Regular images
    ".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif", ".avif",
    # Medical imaging
    ".dcm", ".dicom", ".nii", ".nii.gz",
}

_IMAGE_EXTS   = {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".tiff", ".tif", ".avif"}
_MEDICAL_EXTS = {".dcm", ".dicom", ".nii", ".nii.gz"}


def _resolve_ext(filename: str, data: bytes) -> str:
    """Return the canonical extension, handling .nii.gz and DICOM magic."""
    name = filename.lower()
    if name.endswith(".nii.gz"):
        return ".nii.gz"
    ext = os.path.splitext(name)[1]
    # DICOM files often arrive without extension — detect by preamble magic
    if ext not in _INGEST_ACCEPT and len(data) >= 132 and data[128:132] == b"DICM":
        return ".dcm"
    return ext


def _is_image(ext: str) -> bool:
    return ext in _IMAGE_EXTS or ext in _MEDICAL_EXTS


# ── Vision model helpers ───────────────────────────────────────────────────────

_VISION_MAX_PX     = int(os.getenv("VISION_MAX_PX", "1024"))
_VISION_MAX_TOKENS = int(os.getenv("VISION_MAX_TOKENS", "1200"))
_NO_VISION         = os.getenv("NO_VISION", "0") == "1"


def _available_vision_backends() -> list[str]:
    backends = []
    preferred = VISION_BACKEND
    if preferred:
        backends.append(preferred)
    for backend in ("openai", "anthropic", "ollama"):
        if backend not in backends:
            backends.append(backend)
    return backends


async def _call_vision_backend(backend: str, png_b64: str, prompt: str) -> str:
    if backend == "ollama":
        ollama_vis_model = os.getenv("OLLAMA_VISION_MODEL", OLLAMA_MODEL)
        async with httpx.AsyncClient(timeout=240) as c:
            r = await c.post(
                f"{OLLAMA_BASE_URL}/v1/chat/completions",
                json={
                    "model":      ollama_vis_model,
                    "max_tokens": _VISION_MAX_TOKENS,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {"type": "image_url",
                             "image_url": {"url": f"data:image/png;base64,{png_b64}"}},
                            {"type": "text", "text": prompt},
                        ],
                    }],
                },
            )
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"].strip()

    if backend == "anthropic":
        if not ANTHROPIC_API_KEY:
            raise RuntimeError("Anthropic vision unavailable - no API key configured")
        anthropic_vis_model = os.getenv("ANTHROPIC_VISION_MODEL", os.getenv("VISION_MODEL", ANTHROPIC_CHAT_MODEL))
        async with httpx.AsyncClient(timeout=60) as c:
            r = await c.post(
                "https://api.anthropic.com/v1/messages",
                headers={
                    "x-api-key":         ANTHROPIC_API_KEY,
                    "anthropic-version": "2023-06-01",
                    "Content-Type":      "application/json",
                },
                json={
                    "model":      anthropic_vis_model,
                    "max_tokens": _VISION_MAX_TOKENS,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {"type": "image",
                             "source": {"type": "base64", "media_type": "image/png",
                                        "data": png_b64}},
                            {"type": "text", "text": prompt},
                        ],
                    }],
                },
            )
            r.raise_for_status()
            return r.json()["content"][0]["text"].strip()

    if backend == "openai":
        if not OPENAI_API_KEY:
            raise RuntimeError("OpenAI vision unavailable - no API key configured")
        openai_vis_model = os.getenv("OPENAI_VISION_MODEL", os.getenv("OPENAI_CHAT_MODEL", "gpt-4o"))
        async with httpx.AsyncClient(timeout=60) as c:
            r = await c.post(
                "https://api.openai.com/v1/chat/completions",
                headers={"Authorization": f"Bearer {OPENAI_API_KEY}",
                         "Content-Type": "application/json"},
                json={
                    "model":      openai_vis_model,
                    "max_tokens": _VISION_MAX_TOKENS,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {"type": "image_url",
                             "image_url": {"url": f"data:image/png;base64,{png_b64}",
                                           "detail": "high"}},
                            {"type": "text", "text": prompt},
                        ],
                    }],
                },
            )
            r.raise_for_status()
            return r.json()["choices"][0]["message"]["content"].strip()

    raise RuntimeError(f"Unsupported vision backend: {backend}")


async def _call_vision_api(png_b64: str, prompt: str) -> str:
    """Send a PNG (base64) to the configured vision model and return the description."""
    if _NO_VISION:
        return "[Vision analysis skipped — NO_VISION=1]"

    last_error = None
    for backend in _available_vision_backends():
        try:
            return await _call_vision_backend(backend, png_b64, prompt)
        except Exception as exc:
            last_error = exc
            logger.warning("Vision backend '%s' failed, trying fallback: %r", backend, exc)

    if last_error:
        raise last_error
    return "[Vision unavailable - no backend configured]"


def _pil_to_png_b64(img) -> str:
    from PIL import Image  # noqa: F811
    import base64
    img.thumbnail((_VISION_MAX_PX, _VISION_MAX_PX))
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode()


def _normalize_to_uint8(arr):
    import numpy as np
    arr = arr.astype(np.float32)
    mn, mx = arr.min(), arr.max()
    if mx == mn:
        return np.zeros_like(arr, dtype=np.uint8)
    return ((arr - mn) / (mx - mn) * 255).astype(np.uint8)


def _apply_wl(arr, wc: float, ww: float):
    import numpy as np
    low, high = wc - ww / 2, wc + ww / 2
    return np.clip((arr.astype(np.float32) - low) / (ww) * 255, 0, 255).astype(np.uint8)


def _vision_failure_text(exc: Exception, subject: str) -> str:
    if isinstance(exc, httpx.HTTPStatusError):
        status = exc.response.status_code
        if status == 429:
            return f"[{subject} unavailable - vision provider rate limited (HTTP 429)]"
        return f"[{subject} unavailable - vision provider returned HTTP {status}]"
    detail = str(exc).strip() or exc.__class__.__name__
    return f"[{subject} unavailable - {detail}]"


async def _extract_image_text(filename: str, data: bytes, ext: str) -> tuple[str, dict]:
    """
    Returns (text, image_props) for all image formats.
    text      — descriptive string ready for chunking
    image_props — flat dict merged into the ingest response's image_metadata
    """
    import base64

    # ── DICOM ──────────────────────────────────────────────────────────────────
    if ext in (".dcm", ".dicom"):
        try:
            import pydicom
            from PIL import Image
            import numpy as np

            ds = pydicom.dcmread(io.BytesIO(data), force=True)

            def _tag(attr, default=""):
                v = getattr(ds, attr, default)
                return str(v).strip() if v else default

            def _flt(attr):
                try:
                    v = getattr(ds, attr, None)
                    if v is None:
                        return None
                    if hasattr(v, "__iter__") and not isinstance(v, str):
                        v = list(v)[0]
                    return float(v)
                except Exception:
                    return None

            modality    = _tag("Modality")
            study_desc  = _tag("StudyDescription")
            series_desc = _tag("SeriesDescription")
            institution = _tag("InstitutionName")
            manufacturer = _tag("Manufacturer")
            protocol    = _tag("ProtocolName")
            rows        = _tag("Rows")
            cols        = _tag("Columns")
            bits        = _tag("BitsStored")
            n_frames    = int(getattr(ds, "NumberOfFrames", 1) or 1)
            photometric = _tag("PhotometricInterpretation")
            slice_thick = _tag("SliceThickness")
            pixel_sp    = _tag("PixelSpacing")
            kvp         = _tag("KVP")

            image_props = {
                "image_format": "DICOM",
                "image_modality": modality,
                "image_study_description": study_desc,
                "image_series_description": series_desc,
                "image_institution": institution,
                "image_manufacturer": manufacturer,
                "image_protocol": protocol,
                "image_rows": rows, "image_cols": cols,
                "image_bits": bits, "image_frames": str(n_frames),
                "image_photometric": photometric,
                "image_slice_thickness_mm": slice_thick,
                "image_pixel_spacing_mm": pixel_sp,
                "image_kvp": kvp,
            }

            meta_lines = [
                f"[MEDICAL IMAGE: DICOM{f' / {modality}' if modality else ''}]",
                "",
                "=== DICOM METADATA ===",
                f"Modality            : {modality}" if modality else None,
                f"Study Description   : {study_desc}" if study_desc else None,
                f"Series Description  : {series_desc}" if series_desc else None,
                f"Institution         : {institution}" if institution else None,
                f"Manufacturer        : {manufacturer}" if manufacturer else None,
                f"Protocol            : {protocol}" if protocol else None,
                f"KVP                 : {kvp} kV" if kvp else None,
                f"Slice Thickness     : {slice_thick} mm" if slice_thick else None,
                f"Pixel Spacing       : {pixel_sp} mm" if pixel_sp else None,
                f"Dimensions          : {cols} × {rows} px" if rows and cols else None,
                f"Frames              : {n_frames}" if n_frames > 1 else None,
                f"Bit Depth           : {bits}-bit" if bits else None,
                f"Photometric         : {photometric}" if photometric else None,
            ]
            meta_text = "\n".join(l for l in meta_lines if l is not None)

            # Pixel rendering
            vision_text = ""
            try:
                px = ds.pixel_array  # (rows, cols) or (frames, rows, cols)
                if px.ndim == 3:
                    px = px[px.shape[0] // 2]  # middle frame

                wc = _flt("WindowCenter")
                ww = _flt("WindowWidth")
                gray8 = _apply_wl(px, wc, ww) if (wc is not None and ww and ww > 0) \
                    else _normalize_to_uint8(px)

                # Handle RGB DICOM
                if gray8.ndim == 3:
                    img = Image.fromarray(gray8)
                else:
                    img = Image.fromarray(gray8, mode="L")

                b64 = _pil_to_png_b64(img)
                prompt = (
                    f"This is a medical DICOM image (modality: {modality or 'unknown'}). "
                    "Describe clinically: (1) anatomical region and structures visible, "
                    "(2) notable findings or abnormalities (describe objectively, do not diagnose), "
                    "(3) image orientation and quality, (4) any visible text, annotations, "
                    "measurements, or overlays."
                )
                vision_text = await _call_vision_api(b64, prompt)
            except Exception as e:
                vision_text = _vision_failure_text(e, "Visual analysis")

            text = meta_text
            if vision_text:
                text += f"\n\n=== VISUAL ANALYSIS ===\n{vision_text}"
            return text, {k: v for k, v in image_props.items() if v}

        except ImportError as e:
            raise RuntimeError(
                f"Missing dependency for DICOM: {e}. "
                "Run: pip install pydicom Pillow numpy"
            )

    # ── NIfTI ──────────────────────────────────────────────────────────────────
    if ext in (".nii", ".nii.gz"):
        try:
            import nibabel as nib
            import numpy as np
            import tempfile
            from PIL import Image

            suffix = ".nii.gz" if ext == ".nii.gz" else ".nii"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(data)
                tmp_path = tmp.name
            try:
                nii = nib.load(tmp_path)
                hdr = nii.header
                shape = nii.shape  # (nx, ny, nz[, nt])
                vox   = hdr.get_zooms()
                dtype = str(hdr.get_data_dtype())
                descr = hdr.get("descrip", b"").tobytes().decode("utf-8", errors="replace").strip("\x00").strip()
                arr = nii.get_fdata()
            finally:
                os.unlink(tmp_path)

            nx, ny, nz = (shape + (1, 1))[:3]
            nt = shape[3] if len(shape) > 3 else 1
            dx, dy, dz = (vox + (0, 0, 0))[:3]

            image_props = {
                "image_format": "NIfTI",
                "image_dims": f"{nx}x{ny}x{nz}" + (f"x{nt}" if nt > 1 else ""),
                "image_voxel_size_mm": f"{dx:.3f}x{dy:.3f}x{dz:.3f}",
                "image_dtype": dtype,
                "image_description": descr,
            }

            meta_text = "\n".join(filter(None, [
                "[MEDICAL IMAGE: NIfTI]",
                "",
                "=== NIfTI HEADER ===",
                f"Dimensions   : {nx} × {ny} × {nz}" + (f" × {nt} (time)" if nt > 1 else ""),
                f"Voxel Size   : {dx:.3f} × {dy:.3f} × {dz:.3f} mm",
                f"Data Type    : {dtype}",
                f"Description  : {descr}" if descr else None,
            ]))

            vision_text = ""
            try:
                if arr.ndim >= 3:
                    sl = arr[:, :, arr.shape[2] // 2]
                elif arr.ndim == 2:
                    sl = arr
                else:
                    raise ValueError("Unexpected NIfTI array shape")
                gray8 = _normalize_to_uint8(sl)
                img = Image.fromarray(gray8, mode="L")
                b64 = _pil_to_png_b64(img)
                prompt = (
                    "This is a medical neuroimaging image (NIfTI format). "
                    "Describe: (1) visible brain structures or anatomy, (2) any notable "
                    "features, signal intensities, or abnormalities, (3) image plane and "
                    "orientation, (4) image quality and contrast."
                )
                vision_text = await _call_vision_api(b64, prompt)
            except Exception as e:
                vision_text = _vision_failure_text(e, "Visual analysis")

            text = meta_text
            if vision_text:
                text += f"\n\n=== VISUAL ANALYSIS ===\n{vision_text}"
            return text, {k: v for k, v in image_props.items() if v}

        except ImportError as e:
            raise RuntimeError(
                f"Missing dependency for NIfTI: {e}. "
                "Run: pip install nibabel numpy Pillow"
            )

    # ── Regular images ──────────────────────────────────────────────────────────
    try:
        from PIL import Image
        import base64

        img = Image.open(io.BytesIO(data))
        width, height = img.size
        mode  = img.mode
        fmt   = img.format or ext.lstrip(".")

        image_props = {
            "image_format": fmt.upper(),
            "image_width":  str(width),
            "image_height": str(height),
            "image_mode":   mode,
        }

        # Convert to RGB PNG for vision
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        b64 = _pil_to_png_b64(img)

        prompt = (
            "Describe this image in detail: (1) main subject and overall content, "
            "(2) any text, labels, annotations, or measurements visible, "
            "(3) key visual elements, colors, or patterns, "
            "(4) context or apparent purpose."
        )
        vision_text = await _call_vision_api(b64, prompt)

        text = "\n".join([
            f"[IMAGE: {fmt.upper()}  {width}×{height}px]",
            "",
            vision_text or "[Vision analysis skipped]",
        ])
        return text, image_props

    except ImportError as e:
        raise RuntimeError(f"Missing dependency for images: {e}. Run: pip install Pillow")

def _extract_text(filename: str, data: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(data))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    if ext == ".docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(io.BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    if ext == ".json":
        try:
            return json.dumps(json.loads(data), indent=2)
        except Exception:
            pass
    return data.decode("utf-8", errors="replace")


import re as _re

_ENTITY_PATTERNS = [
    r'\b[A-Z][a-z]+ [A-Z][a-z]+(?:\s[A-Z][a-z]+)?\b',  # Person/place names (Title Case)
    r'\b[A-Z]{2,}\b',                                      # Acronyms (e.g. AIOps, GCOR, AWS)
    r'\b\d{4}-\d{2}-\d{2}\b',                             # ISO dates (2026-03-25)
]


def _find_entity_boundaries(text: str) -> set:
    """Return character positions that are mid-entity — avoid splitting here."""
    bad: set = set()
    for pattern in _ENTITY_PATTERNS:
        for m in _re.finditer(pattern, text):
            bad.update(range(m.start(), m.end()))
    return bad


def _chunk_text(text: str, size: int = 2000, overlap: int = 200) -> list[str]:
    bad_positions = _find_entity_boundaries(text)
    chunks, start = [], 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            for sep in ("\n\n", "\n", " "):
                pos = text.rfind(sep, start + size // 2, end)
                if pos != -1 and pos not in bad_positions:
                    end = pos
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
        if start >= len(text) - overlap:
            break
    return chunks


# The legacy HTML-form /ingest endpoint (and its _neo4j_ingest/_qdrant_ingest
# helpers) has been removed. /api/ingest below is the one ingest entry point;
# it already writes to MinIO (S3) and Graphiti — see _ingest_bytes().

